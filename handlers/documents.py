import os
import json
import anthropic
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes
from telegram.constants import ParseMode
from handlers.texts import t, TEXTS
from handlers.rag_base import get_system_prompt, SELF_EVAL_PROMPT
from database import Database, free_limit_for

# Единый реестр подключаемых документов. Остальные типы продолжают работать
# через существующие DOC_QUESTIONS и общий генератор.
DOCUMENT_REGISTRY = {
    "kindergarten_cycle_schedule": {
        "document_id": "kindergarten_cycle_schedule", "title_ru": "Циклограмма", "title_kz": "Циклограмма", "title_en": "Weekly Cyclogram",
        "category": "kindergarten", "required_fields": ["group", "period", "week_topic"], "optional_fields": ["events"],
        "profile_fields": ["school", "name", "age_group"], "forbidden_facts": ["dates", "events", "children", "results"], "word_structure": "landscape_5x5", "validation_rules": ["required_cycle_fields"], "language_support": ["ru", "kz", "en"]
    },
    "kg_activity_summary": {
        "document_id": "kg_activity_summary", "title_ru": "Технологическая карта ОУД", "title_kz": "ҰОҚ технологиялық картасы", "title_en": "Activity Technological Map",
        "category": "kindergarten", "required_fields": ["topic", "goals", "age_group"], "optional_fields": ["materials"],
        "profile_fields": ["school", "name", "age_group"], "forbidden_facts": ["materials", "children", "results", "methodist_requirements"], "word_structure": "header_3_stages", "validation_rules": ["no_unconfirmed_materials"], "language_support": ["ru", "kz", "en"]
    },
    "kg_individual_development_card": {
        "document_id": "kg_individual_development_card", "title_ru": "Индивидуальная карта развития ребенка", "title_kz": "Баланың жеке даму картасы", "title_en": "Child Individual Development Card",
        "category": "kindergarten", "required_fields": ["child_name", "birth_year_age", "group", "school_year", "observations"], "optional_fields": [],
        "profile_fields": ["school", "age_group"], "forbidden_facts": ["development_levels", "diagnoses", "results"], "word_structure": "five_competency_columns", "validation_rules": ["empty_results_without_observations"], "language_support": ["ru", "kz", "en"]
    },
    "lesson_plan": {
        "document_id": "lesson_plan", "title_ru": "Краткосрочный план (КСП)", "title_kz": "Қысқамерзімді жоспар (ҚМЖ)", "title_en": "Lesson Plan",
        "category": "school", "required_fields": ["subject_class", "topic", "duration"], "optional_fields": ["goals", "date"], "profile_fields": ["school", "name", "subject", "classes"], "forbidden_facts": ["official_goals", "date", "textbook", "resources"], "word_structure": "lesson_table", "validation_rules": ["no_unconfirmed_official_facts"], "language_support": ["ru", "kz", "en"]
    },
    "calendar_plan": {
        "document_id": "calendar_plan", "title_ru": "Календарно-тематический план (КТП)", "title_kz": "Күнтізбелік-тақырыптық жоспар (КТЖ)", "title_en": "Calendar-Thematic Plan",
        "category": "school", "required_fields": ["subject_class", "period", "hours_per_week"], "optional_fields": ["textbook", "dates"], "profile_fields": ["school", "subject", "classes"], "forbidden_facts": ["textbook", "hours", "dates", "official_goals"], "word_structure": "calendar_table", "validation_rules": ["blank_unknown_dates"], "language_support": ["ru", "kz", "en"]
    },
}

REGISTRY_QUESTIONS = {
    "ru": {
        "kg_activity_summary": [{"key": "topic", "q": "Тема занятия?"}, {"key": "age_group", "q": "Группа и возраст?"}, {"key": "goals", "q": "Что дети должны понять или уметь?"}, {"key": "materials", "q": "Какие материалы реально есть? Если нет — напишите «нет»."}],
        "kg_individual_development_card": [{"key": "child_name", "q": "ФИО ребенка?"}, {"key": "birth_year_age", "q": "Год рождения и возраст?"}, {"key": "group", "q": "Группа?"}, {"key": "school_year", "q": "Учебный год?"}, {"key": "observations", "q": "Есть реальные наблюдения или нужен пустой бланк?"}],
        "lesson_plan": [{"key": "subject_class", "q": "Предмет и класс?"}, {"key": "topic", "q": "Тема урока?"}, {"key": "goals", "q": "Цели обучения, если вы их знаете? Если нет — оставим пустыми."}, {"key": "date", "q": "Дата и длительность урока?"}],
        "calendar_plan": [{"key": "subject_class", "q": "Предмет и класс?"}, {"key": "period", "q": "Период?"}, {"key": "textbook", "q": "Учебник или программа? Если нет — напишите «нет»."}, {"key": "hours_per_week", "q": "Часов в неделю? Если неизвестно — напишите «нет»."}, {"key": "dates", "q": "Даты или оставить пустыми?"}],
    },
    "kz": {
        "kg_activity_summary": [{"key": "topic", "q": "ҰОҚ тақырыбы?"}, {"key": "age_group", "q": "Топ және жасы?"}, {"key": "goals", "q": "Балалар нені түсінуі немесе меңгеруі керек?"}, {"key": "materials", "q": "Нақты бар материалдар? Жоқ болса «жоқ» деп жазыңыз."}],
        "kg_individual_development_card": [{"key": "child_name", "q": "Баланың аты-жөні?"}, {"key": "birth_year_age", "q": "Туған жылы және жасы?"}, {"key": "group", "q": "Тобы?"}, {"key": "school_year", "q": "Оқу жылы?"}, {"key": "observations", "q": "Нақты бақылау бар ма, әлде бос үлгі керек пе?"}],
        "lesson_plan": [{"key": "subject_class", "q": "Пән және сынып?"}, {"key": "topic", "q": "Сабақ тақырыбы?"}, {"key": "goals", "q": "Оқу мақсаттары, егер білсеңіз?"}, {"key": "date", "q": "Сабақ күні және ұзақтығы?"}],
        "calendar_plan": [{"key": "subject_class", "q": "Пән және сынып?"}, {"key": "period", "q": "Кезең?"}, {"key": "textbook", "q": "Оқулық немесе бағдарлама?"}, {"key": "hours_per_week", "q": "Аптасына неше сағат?"}, {"key": "dates", "q": "Күндер немесе бос қалдыру керек пе?"}],
    },
    "en": {
        "kg_activity_summary": [{"key": "topic", "q": "Activity topic?"}, {"key": "age_group", "q": "Group and age?"}, {"key": "goals", "q": "What should children understand or learn?"}, {"key": "materials", "q": "Which materials are actually available? If none, write “none”."}],
        "kg_individual_development_card": [{"key": "child_name", "q": "Child’s full name?"}, {"key": "birth_year_age", "q": "Birth year and age?"}, {"key": "group", "q": "Group?"}, {"key": "school_year", "q": "Academic year?"}, {"key": "observations", "q": "Do you have real observations or need a blank form?"}],
        "lesson_plan": [{"key": "subject_class", "q": "Subject and class?"}, {"key": "topic", "q": "Lesson topic?"}, {"key": "goals", "q": "Learning objectives, if known?"}, {"key": "date", "q": "Lesson date and duration?"}],
        "calendar_plan": [{"key": "subject_class", "q": "Subject and class?"}, {"key": "period", "q": "Period?"}, {"key": "textbook", "q": "Textbook or curriculum?"}, {"key": "hours_per_week", "q": "Hours per week?"}, {"key": "dates", "q": "Dates or leave them blank?"}],
    },
}

# ══════════════════════════════════════════════════════════════
# ВОПРОСЫ ДЛЯ ДОКУМЕНТОВ УЧИТЕЛЯ (школа)
# ══════════════════════════════════════════════════════════════
DOC_QUESTIONS = {
    "ru": {
        "lesson_plan": [
            {"key": "subject_class", "q": "📚 Предмет и класс?\n\n_Пример: Математика, 7А_"},
            {"key": "topic",         "q": "📖 Тема урока?"},
            {"key": "duration",      "q": "⏱ Длительность?\n\n_Пример: 45 минут_"},
            {"key": "goals",         "q": "🎯 Цели обучения?\n\n_Или напишите «автоматически»_"},
        ],
        "calendar_plan": [
            {"key": "subject_class",  "q": "📚 Предмет и класс?\n\n_Пример: Алгебра, 8А_"},
            {"key": "month",          "q": "📅 На какой месяц?\n\n_Пример: Ноябрь 2024_"},
            {"key": "hours_per_week", "q": "⏰ Часов в неделю?\n\n_Пример: 3 часа_"},
            {"key": "topics",         "q": "📋 Темы уроков через запятую\n\n_Или напишите «по программе МОН»_"},
        ],
        "lesson_summary": [
            {"key": "subject_class", "q": "📚 Предмет и класс?"},
            {"key": "topic",         "q": "📖 Тема урока?"},
            {"key": "key_points",    "q": "🔑 Ключевые моменты?\n\n_Или напишите «автоматически»_"},
        ],
        "monthly_report": [
            {"key": "period",      "q": "📅 За какой период?\n\n_Пример: Октябрь 2024_"},
            {"key": "classes",     "q": "🏫 Классы?\n\n_Пример: 7А, 8Б, 9В_"},
            {"key": "performance", "q": "📊 Успеваемость?\n\n_Пример: 7А — 65%/90%, 8Б — 70%/95%\n(% качества / % успеваемости)_"},
            {"key": "extra",       "q": "🏆 Внеклассные мероприятия, олимпиады?\n\n_Если нет — напишите «нет»_"},
        ],
        "control_analysis": [
            {"key": "subject_class", "q": "📚 Предмет и класс?"},
            {"key": "date",          "q": "📅 Дата контрольной работы?"},
            {"key": "results",       "q": "📊 Результаты?\n\n_Пример: 5 — 3 уч., 4 — 8 уч., 3 — 5 уч., 2 — 2 уч._"},
            {"key": "topic",         "q": "📖 Тема контрольной работы?"},
        ],
        "sor_soch": [
            {"key": "sor_or_soch",   "q": "📋 СОР (за раздел) или СОЧ (за четверть)?\n\n_Напишите СОР или СОЧ_"},
            {"key": "subject_class", "q": "📚 Предмет и класс?"},
            {"key": "section_topic", "q": "📖 Раздел / тема, по которой оценивание?"},
            {"key": "criteria",      "q": "📊 Критерии оценивания?\n\n_Или напишите «автоматически» — составлю по целям обучения_"},
        ],
        "sor_soch_analysis": [
            {"key": "sor_or_soch",   "q": "📋 Анализ СОР или СОЧ?"},
            {"key": "subject_class", "q": "📚 Предмет и класс?"},
            {"key": "date",          "q": "📅 Дата проведения?"},
            {"key": "results",       "q": "📊 Результаты по баллам?\n\n_Пример: 5уч.-10-12б, 8уч.-7-9б, 4уч.-4-6б, 2уч.-0-3б_"},
            {"key": "typical_errors","q": "⚠️ Типичные ошибки класса?\n\n_Или напишите «автоматически»_"},
        ],
        "characteristic": [
            {"key": "student_name", "q": "👤 ФИО ученика и класс?\n\n_Или выберите из базы ниже_"},
            {"key": "performance",  "q": "📊 Успеваемость?\n\n_Пример: отличник, хорошист, троечник_"},
            {"key": "behavior",     "q": "😊 Поведение и характер?\n\n_Пример: дисциплинированный, активный_"},
            {"key": "activities",   "q": "🏆 Участие в мероприятиях?\n\n_Если нет — напишите «нет»_"},
            {"key": "purpose",      "q": "📋 Цель характеристики?\n\n_Пример: по месту требования_"},
        ],
        "absence_cert": [
            {"key": "student_name",  "q": "👤 ФИО ученика и класс?"},
            {"key": "absence_dates", "q": "📅 Даты отсутствия?\n\n_Пример: 14-16 октября 2024_"},
            {"key": "reason",        "q": "❓ Причина?\n\n_Пример: болезнь (справка есть)_"},
        ],
        "discipline_act": [
            {"key": "student_name", "q": "👤 ФИО ученика и класс?"},
            {"key": "date",         "q": "📅 Дата нарушения?"},
            {"key": "violation",    "q": "⚠️ Описание нарушения?"},
            {"key": "witnesses",    "q": "👥 Свидетели?\n\n_Если нет — напишите «нет»_"},
        ],
        "gratitude_letter": [
            {"key": "student_name", "q": "👤 ФИО ученика и класс?"},
            {"key": "achievement",  "q": "🏆 За что награждается?\n\n_Пример: победа на олимпиаде по математике_"},
        ],
        "parent_letter": [
            {"key": "student_name", "q": "👤 ФИО ученика и класс?"},
            {"key": "topic",        "q": "📋 Тема письма?\n\n_Пример: успеваемость, поведение_"},
            {"key": "details",      "q": "📝 Детали письма?"},
        ],
        "vacation_request": [
            {"key": "vacation_type", "q": "🏖 Вид отпуска?\n\n1️⃣ Ежегодный трудовой (56 дней)\n2️⃣ За свой счёт\n3️⃣ Учебный\n\n_Напишите номер или название_"},
            {"key": "dates",         "q": "📅 Даты?\n\n_Пример: с 01.07.2025 по 25.08.2025_"},
        ],
        "explanation": [
            {"key": "absence_date", "q": "📅 Дата отсутствия или опоздания?\n\n_Пример: 15 октября 2024_"},
            {"key": "reason",       "q": "❓ Причина?\n\n_Пример: болезнь, задержка транспорта_"},
            {"key": "documents",    "q": "📎 Есть подтверждающие документы?\n\n_Если нет — напишите «нет»_"},
        ],
        "announcement": [
            {"key": "event",    "q": "📢 Что за мероприятие?\n\n_Пример: родительское собрание, олимпиада_"},
            {"key": "datetime", "q": "📅 Дата и время?\n\n_Пример: 20 ноября, 18:00_"},
            {"key": "location", "q": "📍 Место проведения?\n\n_Пример: кабинет 205, актовый зал_"},
        ],

        # ── ДОКУМЕНТЫ ДЛЯ САДИКА ──
        "kg_thematic_plan": [
            {"key": "age_group",  "q": "👶 Возрастная группа?\n\n_Пример: средняя группа (4-5 лет)_"},
            {"key": "month",      "q": "📅 На какой месяц?\n\n_Пример: Ноябрь 2024_"},
            {"key": "topics",     "q": "📋 Темы недель через запятую\n\n_Или напишите «по программе»_"},
        ],
        "kg_activity_summary": [
            {"key": "age_group", "q": "👶 Возрастная группа?"},
            {"key": "topic",     "q": "📖 Тема занятия (ОУД)?\n\n_Пример: «Осень золотая» (ознакомление с природой)_"},
            {"key": "goals",     "q": "🎯 Цель занятия?\n\n_Или напишите «автоматически»_"},
        ],
        "kg_perspective_plan": [
            {"key": "age_group", "q": "👶 Возрастная группа?"},
            {"key": "month",     "q": "📅 На какой месяц?\n\n_Пример: Ноябрь 2024_"},
            {"key": "events",    "q": "🎉 Ключевые мероприятия месяца через запятую?\n\n_Или напишите «автоматически»_"},
        ],
        "kg_matinee_script": [
            {"key": "theme",     "q": "🎭 Тема утренника?\n\n_Пример: Новый год, 8 Марта, Наурыз_"},
            {"key": "age_group", "q": "👶 Возрастная группа?"},
            {"key": "date",      "q": "📅 Дата проведения?"},
            {"key": "roles",     "q": "🎬 Роли и участники?\n\n_Пример: Дед Мороз — воспитатель, Снегурочка — муз.руководитель, 5 детей-снежинок_"},
            {"key": "props",     "q": "🎁 Нужный реквизит?\n\n_Или напишите «автоматически»_"},
        ],
        "kg_monthly_report": [
            {"key": "period",    "q": "📅 За какой период?\n\n_Пример: Октябрь 2024_"},
            {"key": "age_group", "q": "👶 Группа?\n\n_Пример: старшая группа «Ромашка»_"},
            {"key": "progress",  "q": "📊 Освоение программы детьми (кратко)?\n\n_Пример: большинство детей усвоили материал по ФЭМП_"},
            {"key": "extra",     "q": "🏆 Мероприятия, утренники, конкурсы?\n\n_Если нет — напишите «нет»_"},
        ],
        "kg_monitoring": [
            {"key": "age_group", "q": "👶 Возрастная группа?"},
            {"key": "period",    "q": "📅 На какой месяц/срез?"},
            {"key": "skills",    "q": "📊 Что осваивали (по образовательным областям)?\n\n_Или напишите «автоматически» — возьму области по ГОСО_"},
            {"key": "results",   "q": "📈 Общий результат по группе?\n\n_Пример: 18 детей — высокий уровень, 6 — средний, 1 — низкий_"},
        ],
        "kg_child_characteristic": [
            {"key": "student_name", "q": "👤 Имя ребёнка, возраст, группа?\n\n_Или выберите из базы ниже_"},
            {"key": "performance",  "q": "📊 Освоение программы?\n\n_Пример: хорошо усваивает материал_"},
            {"key": "behavior",     "q": "😊 Поведение и характер?\n\n_Пример: общительный, любознательный_"},
            {"key": "activities",   "q": "🏆 Достижения, участие в утренниках?\n\n_Если нет — напишите «нет»_"},
            {"key": "purpose",      "q": "📋 Цель характеристики?\n\n_Пример: для психолога, по месту требования_"},
        ],
        "kg_parent_letter": [
            {"key": "student_name", "q": "👤 Имя ребёнка и группа?"},
            {"key": "topic",        "q": "📋 Тема письма?\n\n_Пример: адаптация, поведение, успехи_"},
            {"key": "details",      "q": "📝 Детали письма?"},
        ],
        "kg_absence_cert": [
            {"key": "student_name",  "q": "👤 Имя ребёнка и группа?"},
            {"key": "absence_dates", "q": "📅 Даты отсутствия?"},
            {"key": "reason",        "q": "❓ Причина?\n\n_Пример: болезнь (справка есть)_"},
        ],
        "kg_vacation_request": [
            {"key": "vacation_type", "q": "🏖 Вид отпуска?\n\n1️⃣ Ежегодный трудовой\n2️⃣ За свой счёт\n\n_Напишите номер или название_"},
            {"key": "dates",         "q": "📅 Даты?\n\n_Пример: с 01.07.2025 по 25.08.2025_"},
        ],
        "kg_explanation": [
            {"key": "absence_date", "q": "📅 Дата отсутствия или опоздания?"},
            {"key": "reason",       "q": "❓ Причина?"},
            {"key": "documents",    "q": "📎 Есть подтверждающие документы?\n\n_Если нет — напишите «нет»_"},
        ],
        "kg_announcement": [
            {"key": "event",    "q": "📢 Что за мероприятие?\n\n_Пример: утренник, родительское собрание_"},
            {"key": "datetime", "q": "📅 Дата и время?"},
            {"key": "location", "q": "📍 Место проведения?\n\n_Пример: музыкальный зал_"},
        ],

        # ── ОБЩИЕ ДОКУМЕНТЫ (доступны и школе, и саду) ──
        "parent_work_plan": [
            {"key": "period",  "q": "📅 На какой период план?\n\n_Пример: 2024-2025 учебный год_"},
            {"key": "group",   "q": "🏷 Класс/группа?"},
            {"key": "topics",  "q": "📋 Темы собраний по месяцам?\n\n_Или напишите «автоматически»_"},
        ],
        "upbringing_plan": [
            {"key": "period",  "q": "📅 На какой период план?"},
            {"key": "group",   "q": "🏷 Класс/группа?"},
            {"key": "directions", "q": "📋 Направления работы?\n\n_Пример: патриотическое, ЗОЖ, трудовое. Или «автоматически»_"},
        ],
        "parent_meeting_protocol": [
            {"key": "date",     "q": "📅 Дата собрания?"},
            {"key": "topic",    "q": "📋 Повестка дня?"},
            {"key": "attendees","q": "👥 Сколько человек присутствовало?"},
            {"key": "theses",   "q": "🗣 Кратко о чём говорили выступавшие?\n\n_Или напишите «автоматически»_"},
            {"key": "decision", "q": "✅ Итоговое решение собрания?"},
        ],
        "individual_work_plan": [
            {"key": "student_name", "q": "👤 ФИО/имя ребёнка и класс/группа?\n\n_Или выберите из базы ниже_"},
            {"key": "problem",      "q": "❓ В чём проблема / что нужно подтянуть?\n\n_Пример: отстаёт по чтению_"},
            {"key": "period",       "q": "📅 На какой период рассчитан план?"},
            {"key": "frequency",    "q": "⏰ Периодичность занятий?\n\n_Пример: 2 раза в неделю_"},
        ],
        "housing_survey_act": [
            {"key": "student_name",  "q": "👤 ФИО ребёнка и класс/группа?\n\n_Или выберите из базы ниже_"},
            {"key": "visit_date",    "q": "📅 Дата обследования?"},
            {"key": "commission",    "q": "👥 Состав комиссии?\n\n_Пример: кл.рук. Иванова М.П., соц.педагог Петрова А.Б._"},
            {"key": "family",        "q": "👨‍👩‍👧 Состав семьи?"},
            {"key": "conditions",    "q": "🏠 Условия проживания и рабочее место ребёнка?\n\n_Или напишите «автоматически» на основе состава семьи_"},
        ],
    },
    "kz": {
        "lesson_plan": [
            {"key": "subject_class", "q": "📚 Пән және сынып?\n\n_Мысалы: Математика, 7А_"},
            {"key": "topic",         "q": "📖 Сабақтың тақырыбы?"},
            {"key": "duration",      "q": "⏱ Ұзақтығы?\n\n_Мысалы: 45 минут_"},
            {"key": "goals",         "q": "🎯 Оқу мақсаттары?\n\n_Немесе «автоматты» деп жазыңыз_"},
        ],
        "monthly_report": [
            {"key": "period",      "q": "📅 Қандай кезең?\n\n_Мысалы: Қазан 2024_"},
            {"key": "classes",     "q": "🏫 Сыныптар?\n\n_Мысалы: 7А, 8Б, 9В_"},
            {"key": "performance", "q": "📊 Үлгерім?\n\n_Мысалы: 7А — 65%/90%_"},
            {"key": "extra",       "q": "🏆 Сыныптан тыс іс-шаралар?\n\n_Жоқ болса «жоқ» деп жазыңыз_"},
        ],
        "vacation_request": [
            {"key": "vacation_type", "q": "🏖 Демалыс түрі?\n\n1️⃣ Жылдық еңбек (56 күн)\n2️⃣ Ақысыз\n3️⃣ Оқу\n\n_Нөмірін жазыңыз_"},
            {"key": "dates",         "q": "📅 Күндері?\n\n_Мысалы: 01.07.2025-тен 25.08.2025-ке дейін_"},
        ],
        "explanation": [
            {"key": "absence_date", "q": "📅 Болмаған күн?\n\n_Мысалы: 2024 жылғы 15 қазан_"},
            {"key": "reason",       "q": "❓ Себебі?\n\n_Мысалы: ауырып қалу_"},
            {"key": "documents",    "q": "📎 Растайтын құжаттар?\n\n_Жоқ болса «жоқ»_"},
        ],
        "characteristic": [
            {"key": "student_name", "q": "👤 Оқушының аты-жөні мен сыныбы?"},
            {"key": "performance",  "q": "📊 Үлгерімі?\n\n_Мысалы: үздік, жақсы_"},
            {"key": "behavior",     "q": "😊 Мінез-құлқы?"},
            {"key": "activities",   "q": "🏆 Іс-шараларға қатысуы?\n\n_Немесе «жоқ»_"},
            {"key": "purpose",      "q": "📋 Мінездеме мақсаты?\n\n_Мысалы: талап еткен жерге_"},
        ],
        "announcement": [
            {"key": "event",    "q": "📢 Іс-шара?\n\n_Мысалы: ата-аналар жиналысы_"},
            {"key": "datetime", "q": "📅 Күні мен уақыты?\n\n_Мысалы: 20 қараша, 18:00_"},
            {"key": "location", "q": "📍 Орны?\n\n_Мысалы: 205 кабинет_"},
        ],
        "absence_cert": [
            {"key": "student_name",  "q": "👤 Оқушының аты-жөні мен сыныбы?"},
            {"key": "absence_dates", "q": "📅 Болмаған күндер?"},
            {"key": "reason",        "q": "❓ Себебі?"},
        ],
        "lesson_summary": [
            {"key": "subject_class", "q": "📚 Пән және сынып?"},
            {"key": "topic",         "q": "📖 Тақырып?"},
            {"key": "key_points",    "q": "🔑 Негізгі тұстар немесе «автоматты»?"},
        ],
        "calendar_plan": [
            {"key": "subject_class",  "q": "📚 Пән және сынып?"},
            {"key": "month",          "q": "📅 Қандай ай?\n\n_Мысалы: Қараша 2024_"},
            {"key": "hours_per_week", "q": "⏰ Аптасына неше сағат?"},
            {"key": "topics",         "q": "📋 Тақырыптар немесе «МОН бағдарламасы бойынша»"},
        ],
        "control_analysis": [
            {"key": "subject_class", "q": "📚 Пән және сынып?"},
            {"key": "date",          "q": "📅 Бақылау жұмысының күні?"},
            {"key": "results",       "q": "📊 Нәтижелер?\n\n_Мысалы: 5 — 3 оқ., 4 — 8 оқ._"},
            {"key": "topic",         "q": "📖 Тақырыбы?"},
        ],
        "sor_soch": [
            {"key": "sor_or_soch",   "q": "📋 БЖБ (бөлім) немесе ТЖБ (тоқсан)?"},
            {"key": "subject_class", "q": "📚 Пән және сынып?"},
            {"key": "section_topic", "q": "📖 Бөлім/тақырып?"},
            {"key": "criteria",      "q": "📊 Бағалау критерийлері?\n\n_Немесе «автоматты»_"},
        ],
        "sor_soch_analysis": [
            {"key": "sor_or_soch",   "q": "📋 БЖБ немесе ТЖБ талдауы?"},
            {"key": "subject_class", "q": "📚 Пән және сынып?"},
            {"key": "date",          "q": "📅 Өткізілген күні?"},
            {"key": "results",       "q": "📊 Балл бойынша нәтижелер?"},
            {"key": "typical_errors","q": "⚠️ Сыныптың жиі қателері?\n\n_Немесе «автоматты»_"},
        ],
        "discipline_act": [
            {"key": "student_name", "q": "👤 Оқушының аты-жөні мен сыныбы?"},
            {"key": "date",         "q": "📅 Бұзылған күн?"},
            {"key": "violation",    "q": "⚠️ Бұзушылықтың сипаттамасы?"},
            {"key": "witnesses",    "q": "👥 Куәгерлер?\n\n_Жоқ болса «жоқ»_"},
        ],
        "gratitude_letter": [
            {"key": "student_name", "q": "👤 Оқушының аты-жөні мен сыныбы?"},
            {"key": "achievement",  "q": "🏆 Не үшін марапатталады?"},
        ],
        "parent_letter": [
            {"key": "student_name", "q": "👤 Оқушының аты-жөні мен сыныбы?"},
            {"key": "topic",        "q": "📋 Хат тақырыбы?"},
            {"key": "details",      "q": "📝 Мәліметтер?"},
        ],

        # ── САДИК (KZ) ──
        "kg_thematic_plan": [
            {"key": "age_group", "q": "👶 Жас тобы?"},
            {"key": "month",     "q": "📅 Қандай ай?"},
            {"key": "topics",    "q": "📋 Апталық тақырыптар немесе «бағдарлама бойынша»"},
        ],
        "kg_activity_summary": [
            {"key": "age_group", "q": "👶 Жас тобы?"},
            {"key": "topic",     "q": "📖 Сабақтың тақырыбы?"},
            {"key": "goals",     "q": "🎯 Сабақтың мақсаты немесе «автоматты»?"},
        ],
        "kg_perspective_plan": [
            {"key": "age_group", "q": "👶 Жас тобы?"},
            {"key": "month",     "q": "📅 Қандай ай?"},
            {"key": "events",    "q": "🎉 Айдың негізгі іс-шаралары?\n\n_Немесе «автоматты»_"},
        ],
        "kg_matinee_script": [
            {"key": "theme",     "q": "🎭 Мерекенің тақырыбы?\n\n_Мысалы: Жаңа жыл, Наурыз_"},
            {"key": "age_group", "q": "👶 Жас тобы?"},
            {"key": "date",      "q": "📅 Күні?"},
            {"key": "roles",     "q": "🎬 Рөлдер мен қатысушылар?"},
            {"key": "props",     "q": "🎁 Керекті реквизит?\n\n_Немесе «автоматты»_"},
        ],
        "kg_monthly_report": [
            {"key": "period",    "q": "📅 Қандай кезең?"},
            {"key": "age_group", "q": "👶 Топ?"},
            {"key": "progress",  "q": "📊 Балалардың бағдарламаны меңгеруі?"},
            {"key": "extra",     "q": "🏆 Іс-шаралар, мерекелер?\n\n_Жоқ болса «жоқ»_"},
        ],
        "kg_monitoring": [
            {"key": "age_group", "q": "👶 Жас тобы?"},
            {"key": "period",    "q": "📅 Қандай ай/кезең?"},
            {"key": "skills",    "q": "📊 Қандай дағдылар бойынша?\n\n_Немесе «автоматты»_"},
            {"key": "results",   "q": "📈 Топ бойынша нәтиже?"},
        ],
        "kg_child_characteristic": [
            {"key": "student_name", "q": "👤 Баланың аты, жасы, тобы?"},
            {"key": "performance",  "q": "📊 Бағдарламаны меңгеруі?"},
            {"key": "behavior",     "q": "😊 Мінез-құлқы?"},
            {"key": "activities",   "q": "🏆 Жетістіктері?\n\n_Жоқ болса «жоқ»_"},
            {"key": "purpose",      "q": "📋 Мінездеме мақсаты?"},
        ],
        "kg_parent_letter": [
            {"key": "student_name", "q": "👤 Баланың аты және тобы?"},
            {"key": "topic",        "q": "📋 Хат тақырыбы?"},
            {"key": "details",      "q": "📝 Мәліметтер?"},
        ],
        "kg_absence_cert": [
            {"key": "student_name",  "q": "👤 Баланың аты және тобы?"},
            {"key": "absence_dates", "q": "📅 Болмаған күндер?"},
            {"key": "reason",        "q": "❓ Себебі?"},
        ],
        "kg_vacation_request": [
            {"key": "vacation_type", "q": "🏖 Демалыс түрі?"},
            {"key": "dates",         "q": "📅 Күндері?"},
        ],
        "kg_explanation": [
            {"key": "absence_date", "q": "📅 Болмаған күн?"},
            {"key": "reason",       "q": "❓ Себебі?"},
            {"key": "documents",    "q": "📎 Растайтын құжаттар?\n\n_Жоқ болса «жоқ»_"},
        ],
        "kg_announcement": [
            {"key": "event",    "q": "📢 Іс-шара?\n\n_Мысалы: мереке, ата-аналар жиналысы_"},
            {"key": "datetime", "q": "📅 Күні мен уақыты?"},
            {"key": "location", "q": "📍 Орны?"},
        ],

        # ── ЖАЛПЫ ҚҰЖАТТАР (мектеп пен балабақшаға ортақ) ──
        "parent_work_plan": [
            {"key": "period", "q": "📅 Қандай кезеңге жоспар?"},
            {"key": "group",  "q": "🏷 Сынып/топ?"},
            {"key": "topics", "q": "📋 Айлар бойынша жиналыс тақырыптары?\n\n_Немесе «автоматты»_"},
        ],
        "upbringing_plan": [
            {"key": "period",     "q": "📅 Қандай кезеңге жоспар?"},
            {"key": "group",      "q": "🏷 Сынып/топ?"},
            {"key": "directions", "q": "📋 Жұмыс бағыттары?\n\n_Немесе «автоматты»_"},
        ],
        "parent_meeting_protocol": [
            {"key": "date",      "q": "📅 Жиналыс күні?"},
            {"key": "topic",     "q": "📋 Күн тәртібі?"},
            {"key": "attendees", "q": "👥 Қанша адам қатысты?"},
            {"key": "theses",    "q": "🗣 Сөйлеушілер не туралы айтты?\n\n_Немесе «автоматты»_"},
            {"key": "decision",  "q": "✅ Жиналыстың қорытынды шешімі?"},
        ],
        "individual_work_plan": [
            {"key": "student_name", "q": "👤 Баланың аты-жөні және сынып/тобы?"},
            {"key": "problem",      "q": "❓ Қандай мәселе бар?"},
            {"key": "period",       "q": "📅 Қандай кезеңге жоспар?"},
            {"key": "frequency",    "q": "⏰ Сабақтардың жиілігі?"},
        ],
        "housing_survey_act": [
            {"key": "student_name", "q": "👤 Баланың аты-жөні және сынып/тобы?"},
            {"key": "visit_date",   "q": "📅 Тексеру күні?"},
            {"key": "commission",   "q": "👥 Комиссия құрамы?"},
            {"key": "family",       "q": "👨‍👩‍👧 Отбасы құрамы?"},
            {"key": "conditions",   "q": "🏠 Тұрғын үй жағдайы?\n\n_Немесе «автоматты»_"},
        ],
    },
    "en": {
        "lesson_plan": [
            {"key": "subject_class", "q": "📚 Subject and class?\n\n_Example: Mathematics, Grade 7A_"},
            {"key": "topic",         "q": "📖 Lesson topic?"},
            {"key": "duration",      "q": "⏱ Duration?\n\n_Example: 45 minutes_"},
            {"key": "goals",         "q": "🎯 Learning objectives?\n\n_Or write «automatic»_"},
        ],
        "vacation_request": [
            {"key": "vacation_type", "q": "🏖 Type of leave?\n\n1️⃣ Annual leave\n2️⃣ Unpaid leave\n3️⃣ Study leave"},
            {"key": "dates",         "q": "📅 Dates?\n\n_Example: from 01.07.2025 to 25.08.2025_"},
        ],
        "explanation": [
            {"key": "absence_date", "q": "📅 Date of absence?\n\n_Example: October 15, 2024_"},
            {"key": "reason",       "q": "❓ Reason?\n\n_Example: illness, transport delay_"},
            {"key": "documents",    "q": "📎 Supporting documents?\n\n_If none — write «none»_"},
        ],
        "announcement": [
            {"key": "event",    "q": "📢 Event?\n\n_Example: parent meeting, olympiad_"},
            {"key": "datetime", "q": "📅 Date and time?\n\n_Example: November 20, 6:00 PM_"},
            {"key": "location", "q": "📍 Location?\n\n_Example: room 205, assembly hall_"},
        ],
    }
}

DOC_NAMES = {
    "ru": {
        "lesson_plan": "Краткосрочный план (КСП)",
        "calendar_plan": "Календарно-тематический план (КТП)",
        "lesson_summary": "Конспект урока",
        "monthly_report": "Отчёт учителя",
        "control_analysis": "Анализ контрольной работы",
        "sor_soch": "СОР / СОЧ",
        "sor_soch_analysis": "Анализ СОР / СОЧ",
        "characteristic": "Характеристика ученика",
        "absence_cert": "Справка об отсутствии",
        "discipline_act": "Акт о нарушении дисциплины",
        "gratitude_letter": "Благодарственное письмо",
        "parent_letter": "Письмо родителям",
        "vacation_request": "Заявление на отпуск",
        "explanation": "Объяснительная записка",
        "announcement": "Объявление",
        # садик
        "kg_thematic_plan": "Тематический план занятий",
        "kg_activity_summary": "Технологическая карта ОУД",
        "kg_individual_development_card": "Индивидуальная карта развития ребенка",
        "kindergarten_cycle_schedule": "Циклограмма",
        "development_monitoring": "Мониторинг развития",
        "kg_perspective_plan": "Перспективный план работы",
        "kg_matinee_script": "Сценарий утренника",
        "kg_monthly_report": "Отчёт воспитателя",
        "kg_monitoring": "Мониторинг развития",
        "kg_child_characteristic": "Характеристика воспитанника",
        "kg_parent_letter": "Письмо родителям",
        "kg_absence_cert": "Справка об отсутствии ребёнка",
        "kg_vacation_request": "Заявление на отпуск",
        "kg_explanation": "Объяснительная записка",
        "kg_announcement": "Объявление для родителей",
        # общие
        "parent_work_plan": "План работы с родителями",
        "upbringing_plan": "План воспитательной работы",
        "parent_meeting_protocol": "Протокол родительского собрания",
        "individual_work_plan": "Индивидуальный план работы",
        "housing_survey_act": "Акт обследования жилищно-бытовых условий",
    },
    "kz": {
        "lesson_plan": "Қысқамерзімді жоспар (ҚМЖ)",
        "calendar_plan": "Күнтізбелік-тақырыптық жоспар (КТЖ)",
        "lesson_summary": "Сабақ конспектісі",
        "monthly_report": "Мұғалім есебі",
        "control_analysis": "Бақылау жұмысын талдау",
        "sor_soch": "БЖБ / ТЖБ",
        "sor_soch_analysis": "БЖБ / ТЖБ талдауы",
        "characteristic": "Оқушы мінездемесі",
        "absence_cert": "Болмағаны туралы анықтама",
        "discipline_act": "Тәртіп бұзу актісі",
        "gratitude_letter": "Алғыс хат",
        "parent_letter": "Ата-аналарға хат",
        "vacation_request": "Демалыс өтініші",
        "explanation": "Түсіндірме хат",
        "announcement": "Хабарландыру",
        # балабақша
        "kg_thematic_plan": "Тақырыптық жоспар",
        "kg_activity_summary": "ҰОҚ технологиялық картасы",
        "kg_individual_development_card": "Баланың жеке даму картасы",
        "kindergarten_cycle_schedule": "Циклограмма",
        "development_monitoring": "Даму мониторингі",
        "kg_perspective_plan": "Перспективалық жұмыс жоспары",
        "kg_matinee_script": "Мереке сценарийі",
        "kg_monthly_report": "Тәрбиеші есебі",
        "kg_monitoring": "Даму мониторингі",
        "kg_child_characteristic": "Тәрбиеленуші мінездемесі",
        "kg_parent_letter": "Ата-аналарға хат",
        "kg_absence_cert": "Баланың болмағаны туралы анықтама",
        "kg_vacation_request": "Демалыс өтініші",
        "kg_explanation": "Түсіндірме хат",
        "kg_announcement": "Ата-аналарға хабарландыру",
        # ортақ
        "parent_work_plan": "Ата-аналармен жұмыс жоспары",
        "upbringing_plan": "Тәрбие жұмысының жоспары",
        "parent_meeting_protocol": "Ата-аналар жиналысының хаттамасы",
        "individual_work_plan": "Жеке жұмыс жоспары",
        "housing_survey_act": "Тұрмыстық жағдайды тексеру актісі",
    },
    "en": {
        "lesson_plan": "Lesson Plan",
        "calendar_plan": "Calendar-Thematic Plan",
        "lesson_summary": "Lesson Summary",
        "monthly_report": "Teacher Report",
        "control_analysis": "Test Analysis",
        "sor_soch": "Summative Assessment",
        "sor_soch_analysis": "Summative Assessment Analysis",
        "characteristic": "Student Reference",
        "absence_cert": "Absence Certificate",
        "discipline_act": "Disciplinary Act",
        "gratitude_letter": "Gratitude Letter",
        "parent_letter": "Letter to Parents",
        "vacation_request": "Leave Application",
        "explanation": "Explanatory Note",
        "announcement": "Announcement",
        "kg_thematic_plan": "Thematic Activity Plan",
        "kg_activity_summary": "Activity Technological Map",
        "kg_individual_development_card": "Child Individual Development Card",
        "kindergarten_cycle_schedule": "Weekly Cyclogram",
        "kg_perspective_plan": "Monthly Perspective Plan",
        "kg_matinee_script": "Matinee Script",
        "kg_monthly_report": "Kindergarten Teacher Report",
        "kg_monitoring": "Development Monitoring",
        "kg_child_characteristic": "Child Reference",
        "kg_parent_letter": "Letter to Parents",
        "kg_absence_cert": "Child Absence Certificate",
        "kg_vacation_request": "Leave Application",
        "kg_explanation": "Explanatory Note",
        "kg_announcement": "Announcement for Parents",
        "parent_work_plan": "Parent Engagement Plan",
        "upbringing_plan": "Upbringing Work Plan",
        "parent_meeting_protocol": "Parent Meeting Protocol",
        "individual_work_plan": "Individual Work Plan",
        "housing_survey_act": "Living Conditions Survey Act",
    }
}

# Категории документов учителя (школа)
CAT_DOCS = {
    "planning": ["lesson_plan", "calendar_plan", "lesson_summary"],
    "reports":  ["monthly_report", "control_analysis", "sor_soch", "sor_soch_analysis"],
    "students": ["characteristic", "absence_cert", "discipline_act", "gratitude_letter", "parent_letter"],
    "personal": ["vacation_request", "explanation", "announcement"],
}

# Категории документов садика (полностью отдельный набор — НЕ школьные типы)
CAT_DOCS_KG = {
    "kg_planning": ["kg_thematic_plan", "kg_activity_summary", "kg_individual_development_card", "kindergarten_cycle_schedule", "kg_perspective_plan", "kg_matinee_script"],
    "kg_reports":  ["kg_monthly_report", "development_monitoring"],
    "kg_children": ["kg_child_characteristic", "kg_parent_letter", "kg_absence_cert"],
    "kg_personal": ["kg_vacation_request", "kg_explanation", "kg_announcement"],
}

# Документы, общие для школы и садика — показываются В ОБЕИХ ролях в отдельной категории.
# Каждый тип определён ОДИН раз (без kg_-префикса): нужная терминология (ученик/класс
# или воспитанник/группа) подставляется автоматически системным промптом по роли —
# см. rag_base.get_system_prompt (COMMON_DOC_RULES добавляется в обе ветки).
CAT_DOCS_COMMON = {
    "common": [
        "parent_work_plan", "upbringing_plan", "parent_meeting_protocol",
        "individual_work_plan", "housing_survey_act",
    ],
}

# Объединённый словарь категорий — используется при поиске документов по категории,
# независимо от роли (роль определяет только какое меню категорий показать)
CAT_DOCS_ALL = {**CAT_DOCS, **CAT_DOCS_KG, **CAT_DOCS_COMMON}

# Типы документов, для которых предлагается выбор ребёнка/ученика из базы
STUDENT_LINKED_DOC_TYPES = [
    "characteristic", "absence_cert", "discipline_act", "gratitude_letter", "parent_letter",
    "kg_child_characteristic", "kg_parent_letter", "kg_absence_cert",
    "individual_work_plan", "housing_survey_act",
]

# Отдельный, минимальный сценарий для циклограммы сада. Остальные документы
# пока продолжают использовать существующий реестр вопросов без изменений.
KINDERGARTEN_CYCLE_SCHEDULE = "kindergarten_cycle_schedule"
DEVELOPMENT_MONITORING = "development_monitoring"
CYCLE_REQUIRED_FIELDS = {
    "group": "группа",
    "period": "неделя или даты",
    "week_topic": "тема недели",
}


def validate_cycle_schedule_answers(answers: dict) -> list[str]:
    """Возвращает названия обязательных незаполненных полей циклограммы."""
    return [label for key, label in CYCLE_REQUIRED_FIELDS.items() if not str(answers.get(key, "")).strip()]

# Красивые разделители для дизайна
DIVIDER = "─" * 20

def _build_profile_context(user: dict, lang: str) -> str:
    """Собирает весь профиль пользователя для автоподстановки в промпт."""
    role = user.get("role", "teacher")
    if role == "kindergarten":
        lines = [
            "ПРОФИЛЬ ВОСПИТАТЕЛЯ (используй автоматически):",
            f"- ФИО: {user.get('name') or '[не указано]'}",
            f"- Детский сад: {user.get('school') or '[не указано]'}",
            f"- Возрастная группа: {user.get('age_group') or '[не указано]'}",
            f"- Должность: {user.get('position') or 'воспитатель'}",
            f"- Заведующая: {user.get('director') or '[не указано]'}",
        ]
    else:
        is_ct = "Да" if user.get("is_class_teacher") else "Нет"
        lines = [
            "ПРОФИЛЬ УЧИТЕЛЯ (используй автоматически):",
            f"- ФИО: {user.get('name') or '[не указано]'}",
            f"- Школа: {user.get('school') or '[не указано]'}",
            f"- Должность: {user.get('position') or 'учитель'}",
            f"- Предмет: {user.get('subject') or '[не указано]'}",
            f"- Классы: {user.get('classes') or '[не указано]'}",
            f"- Классный руководитель: {is_ct}",
            f"- Директор: {user.get('director') or '[не указано]'}",
        ]
    return "\n".join(lines)

def _progress_bar(current, total):
    filled = "█" * current
    empty  = "░" * (total - current)
    return f"{filled}{empty} {current}/{total}"

def _doc_lang_name(doc_lang):
    return {"ru": "🇷🇺 Русский", "kz": "🇰🇿 Қазақша", "en": "🇬🇧 English"}.get(doc_lang, doc_lang)


class DocumentHandler:
    def __init__(self, db: Database, api_key: str):
        self.db      = db
        self.api_key = api_key

    async def callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        query   = update.callback_query
        await query.answer()
        data    = query.data
        user_id = update.effective_user.id
        user    = await self.db.get_user(user_id)
        lang    = user.get("lang", "ru") if user else "ru"

        # ── Отмена генерации ──
        if data == "doc_cancel":
            context.user_data.clear()
            from handlers.main_menu import MainMenuHandler
            await query.edit_message_text(
                "❌ Отменено." if lang == "ru" else "❌ Болдырылмады."
            )
            await MainMenuHandler(self.db)._send_main_menu(
                query.message.chat_id, context, user_id, lang
            )
            return

        if data == "doc_template":
            if not user or user.get("role") != "kindergarten":
                await query.edit_message_text(
                    "📄 Загрузка образцов доступна только воспитателям детского сада."
                    if lang == "ru" else
                    "📄 Үлгі жүктеу тек балабақша тәрбиешілеріне қолжетімді."
                )
                return
            labels = ("Циклограмма", "Мониторинг развития", "Отмена") if lang == "ru" else ("Циклограмма", "Даму мониторингі", "Бас тарту")
            await query.edit_message_text(
                "Для какого документа загружаешь образец?" if lang == "ru" else "Үлгіні қай құжат үшін жүктейсіз?",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton(labels[0], callback_data="doc_template_kindergarten_cycle_schedule")],
                    [InlineKeyboardButton(labels[1], callback_data="doc_template_development_monitoring")],
                    [InlineKeyboardButton(labels[2], callback_data="doc_cancel")],
                ])
            )
            return
        if data.startswith("doc_template_"):
            template_type = data[len("doc_template_"):]
            if template_type in {KINDERGARTEN_CYCLE_SCHEDULE, DEVELOPMENT_MONITORING}:
                if not user or user.get("role") != "kindergarten":
                    await query.edit_message_text(
                        "📄 Загрузка образцов доступна только воспитателям детского сада."
                        if lang == "ru" else
                        "📄 Үлгі жүктеу тек балабақша тәрбиешілеріне қолжетімді."
                    )
                    return
                context.user_data["template_doc_type"] = template_type
                context.user_data["step"] = "template_upload"
                await query.edit_message_text(
                    "Отправь Word-файл .docx, который принимает твой методист" if lang == "ru" else "Әдіскер қабылдайтын Word .docx файлын жіберіңіз",
                    reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("Отмена" if lang == "ru" else "Бас тарту", callback_data="doc_cancel")]])
                )
                return
        if data == "doc_template_personal":
            pending = context.user_data.get("template_pending")
            if pending:
                await self.db.save_user_template(user_id, pending["doc_type"], pending["path"], pending["name"], lang, pending["metadata"])
                context.user_data.clear()
                await query.edit_message_text(
                    "Образец сохранен. Точное повторение оформления будет доступно после проверки шаблона" if lang == "ru" else "Үлгі сақталды. Безендіруді дәл қайталау үлгі тексерілгеннен кейін қолжетімді болады"
                )
            return
        if data == "doc_template_org":
            await query.edit_message_text("Пока общий шаблон может добавить только администратор организации" if lang == "ru" else "Әзірге ортақ үлгіні тек ұйым әкімшісі қоса алады")
            return
        if data.startswith("doc_template_delete_") and data != "doc_template_delete_confirm":
            dtype = data[len("doc_template_delete_"):]
            context.user_data["delete_template_type"] = dtype
            await query.edit_message_text(
                "Удалить личный образец?" if lang == "ru" else "Жеке үлгіні жою керек пе?",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("Да, удалить" if lang == "ru" else "Иә, жою", callback_data="doc_template_delete_confirm")], [InlineKeyboardButton("Отмена" if lang == "ru" else "Бас тарту", callback_data="doc_cancel")]])
            )
            return
        if data == "doc_template_delete_confirm":
            path = await self.db.delete_user_template(user_id, context.user_data.get("delete_template_type", ""))
            if path and os.path.exists(path): os.remove(path)
            context.user_data.clear()
            await query.edit_message_text("Личный образец удален" if lang == "ru" else "Жеке үлгі жойылды")
            return
        if data == "doc_template_manage":
            kb = []
            for dtype, ru, kz in [(KINDERGARTEN_CYCLE_SCHEDULE, "Удалить мой образец циклограммы", "Циклограмма үлгісін жою"), (DEVELOPMENT_MONITORING, "Удалить мой образец мониторинга", "Мониторинг үлгісін жою")]:
                if await self.db.get_user_template(user_id, dtype):
                    kb.append([InlineKeyboardButton(ru if lang == "ru" else kz, callback_data=f"doc_template_delete_{dtype}")])
            kb.append([InlineKeyboardButton("Отмена" if lang == "ru" else "Бас тарту", callback_data="doc_cancel")])
            await query.edit_message_text("Личные образцы" if lang == "ru" else "Жеке үлгілер", reply_markup=InlineKeyboardMarkup(kb))
            return
        if data == "doc_monitoring_data_yes":
            context.user_data["step"] = "monitoring_data"
            await query.edit_message_text("Отправь данные строками: ФИО | физическое | коммуникативное | познавательное | творческое | социально-эмоциональное | примечание" if lang == "ru" else "Деректерді жолдармен жіберіңіз: Аты-жөні | дене | коммуникативтік | танымдық | шығармашылық | әлеуметтік-эмоциялық | ескертпе")
            return
        if data == "doc_monitoring_data_no":
            context.user_data["doc_answers"]["rows"] = []
            await self._generate_monitoring(query.message, context, lang)
            return

        # ── Выбор языка документа ──
        if data.startswith("doc_lang_"):
            doc_lang = data.split("_")[2]
            context.user_data["doc_lang"] = doc_lang
            doc_type = context.user_data.get("doc_type", "")
            if doc_type == KINDERGARTEN_CYCLE_SCHEDULE:
                await self._start_cycle_schedule(query, context, user, lang)
                return
            q_lang   = doc_lang if doc_lang in DOC_QUESTIONS else "ru"
            qs = REGISTRY_QUESTIONS.get(doc_lang, {}).get(doc_type) or DOC_QUESTIONS.get(q_lang, DOC_QUESTIONS["ru"]).get(doc_type, [])
            if not qs:
                fallback_q = {
                    "ru": "✍️ Опишите подробно что нужно создать:",
                    "kz": "✍️ Не жасау керектігін сипаттаңыз:",
                    "en": "✍️ Describe what you need in detail:",
                }
                qs = [{"key": "description", "q": fallback_q.get(doc_lang, "✍️ Describe:")}]
            profile_values = {
                "age_group": user.get("age_group", ""), "group": user.get("age_group", ""),
                "subject_class": ", ".join(x for x in [user.get("subject", ""), user.get("classes", "")] if x),
            }
            answers = context.user_data.setdefault("doc_answers", {})
            for key, value in profile_values.items():
                if value and key in {q["key"] for q in qs}:
                    answers[key] = value
            context.user_data["questions"] = qs
            context.user_data["step"]      = "waiting_answer"
            doc_name = DOC_NAMES.get(doc_lang, DOC_NAMES["ru"]).get(doc_type, doc_type)
            await self._ask_question(query.message, context, lang, 0, edit=True, query=query, doc_name=doc_name)
            return

        if data == "doc_lang_confirm":
            doc_lang = context.user_data.get("doc_lang", lang)
            doc_type = context.user_data.get("doc_type", "")
            if doc_type == KINDERGARTEN_CYCLE_SCHEDULE:
                await self._start_cycle_schedule(query, context, user, lang)
                return
            qs = REGISTRY_QUESTIONS.get(doc_lang, {}).get(doc_type) or DOC_QUESTIONS.get(doc_lang, DOC_QUESTIONS["ru"]).get(doc_type, [])
            context.user_data["questions"] = qs or [{"key": "description", "q": "✍️ Опишите, что нужно создать:"}]
            context.user_data["step"] = "waiting_answer"
            await self._ask_question(query.message, context, lang, 0, edit=True, query=query, doc_name=DOC_NAMES.get(doc_lang, DOC_NAMES["ru"]).get(doc_type, doc_type))
            return

        if data == "doc_lang_change":
            await query.edit_message_text(
                "Выберите язык документа / Құжат тілін таңдаңыз / Choose document language",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("🇰🇿 Қазақша", callback_data="doc_lang_kz"), InlineKeyboardButton("🇷🇺 Русский", callback_data="doc_lang_ru")],
                    [InlineKeyboardButton("🇬🇧 English", callback_data="doc_lang_en")],
                ])
            )
            return

        # ── Категория документов (и школьная, и садиковская — cat_planning / cat_kg_planning и т.д.) ──
        if data.startswith("cat_"):
            cat = data[len("cat_"):]
            await self._show_doc_list(query, lang, cat)

        # ── Выбор конкретного документа ──
        elif data.startswith("doc_"):
            doc_type = data[4:]
            await self._start_doc(query, context, user_id, user, lang, doc_type)

        # ── Помощь с текстом ──
        elif data.startswith("ans_help_"):
            field = data[9:]
            context.user_data["help_field"] = field
            context.user_data["step"]       = "help_write"
            doc_type = context.user_data.get("doc_type", "")
            text = (
                f"✍️ *Помощь с текстом*\n\n"
                f"Напишите кратко что хотите сказать — я оформлю официально:"
            )
            kb = [[InlineKeyboardButton("❌ Отмена", callback_data="menu_main")]]
            await query.edit_message_text(text, reply_markup=InlineKeyboardMarkup(kb), parse_mode=ParseMode.MARKDOWN)

        # ── Выбрать ученика/воспитанника из базы ──
        elif data.startswith("gen_student_"):
            student_id = int(data[12:])
            await self._use_student_data(query, context, user_id, user, lang, student_id)

    async def _show_doc_list(self, query, lang, cat):
        docs  = CAT_DOCS_ALL.get(cat, [])
        names = DOC_NAMES.get(lang, DOC_NAMES["ru"])
        keyboard = [[InlineKeyboardButton(names.get(d, d), callback_data=f"doc_{d}")] for d in docs]
        if cat.startswith("kg_"):
            keyboard.append([InlineKeyboardButton("📄 Загрузить образец" if lang == "ru" else "📄 Үлгіні жүктеу", callback_data="doc_template")])
            keyboard.append([InlineKeyboardButton("Личные образцы" if lang == "ru" else "Жеке үлгілер", callback_data="doc_template_manage")])
        keyboard.append([InlineKeyboardButton("◀️ " + t(lang, "back"), callback_data="menu_create")])

        cat_name = t(lang, f"cat_{cat}")
        header = "📂 *{cat}*\n{div}\nВыберите тип документа:" if lang == "ru" else "📂 *{cat}*\n{div}\nҚұжат түрін таңдаңыз:"

        if not docs:
            # Защита от несуществующей/незаполненной категории — не показываем пустой экран молча
            empty_text = (
                "⚠️ Для этой категории пока нет документов." if lang == "ru"
                else "⚠️ Бұл санатта әзірге құжат жоқ."
            )
            await query.edit_message_text(
                empty_text,
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("◀️ " + t(lang, "back"), callback_data="menu_create")]])
            )
            return

        await query.edit_message_text(
            header.format(cat=cat_name, div=DIVIDER),
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode=ParseMode.MARKDOWN
        )

    async def _start_doc(self, query, context, user_id, user, lang, doc_type):
        subscribed = user.get("subscribed", 0)
        free_used  = user.get("free_used", 0)

        if not subscribed and free_used >= free_limit_for(user):
            from handlers.profile import TIER_PRICES
            promo_available = not user.get("promo_used")

            if promo_available:
                text = (
                    f"🚫 *Бесплатные попытки закончились.*\n\n"
                    f"🔥 Месяц PRO с безлимитной генерацией — всего *{TIER_PRICES['pro_promo']} тг* "
                    f"(вместо {TIER_PRICES['pro']} тг)!\n"
                    f"Это разовое предложение, доступно один раз для нового пользователя.\n\n"
                    f"После оплаты пришлите чек — доступ откроется автоматически."
                ) if lang == "ru" else (
                    f"🚫 *Тегін әрекеттер аяқталды.*\n\n"
                    f"🔥 Шексіз PRO — небары *{TIER_PRICES['pro_promo']} тг* ({TIER_PRICES['pro']} тг орнына)!\n"
                    f"Бұл жаңа пайдаланушыға бір реттік ұсыныс.\n\n"
                    f"Төлемнен кейін чекті жіберіңіз — қолжетімділік автоматты ашылады."
                )
                keyboard = [
                    [InlineKeyboardButton(
                        f"🔥 Активировать PRO за {TIER_PRICES['pro_promo']} тг" if lang == "ru" else f"🔥 PRO-ды {TIER_PRICES['pro_promo']} тг-ге белсендіру",
                        callback_data="prof_choose_pro_promo"
                    )],
                    [InlineKeyboardButton("🎁 Или пригласить и получить +5" if lang == "ru" else "🎁 Немесе шақырып +5 алу", callback_data="menu_invite")],
                    [InlineKeyboardButton("Все тарифы" if lang == "ru" else "Барлық тарифтер", callback_data="prof_sub")],
                    [InlineKeyboardButton("◀️ " + t(lang, "back"), callback_data="menu_main")],
                ]
            else:
                text = t(lang, "limit_reached")
                keyboard = [
                    [InlineKeyboardButton("⭐ Оформить подписку", callback_data="prof_sub")],
                    [InlineKeyboardButton("◀️ " + t(lang, "back"), callback_data="menu_main")],
                ]

            await query.edit_message_text(
                text,
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode=ParseMode.MARKDOWN
            )
            return

        # Выбор языка всегда относится только к текущему документу. Не используем
        # язык из профиля и не оставляем его от предыдущей генерации.
        context.user_data.pop("doc_lang", None)
        context.user_data["doc_type"] = doc_type
        doc_name = DOC_NAMES.get(lang, DOC_NAMES["ru"]).get(doc_type, doc_type)
        await query.edit_message_text(
            (f"📄 *{doc_name}*\n{DIVIDER}\n🌐 На каком языке создать документ?" if lang == "ru"
             else f"📄 *{doc_name}*\n{DIVIDER}\n🌐 Құжатты қандай тілде жасау керек?"),
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("🇷🇺 Русский", callback_data="doc_lang_ru"), InlineKeyboardButton("🇰🇿 Қазақша", callback_data="doc_lang_kz")],
                [InlineKeyboardButton("🇬🇧 English", callback_data="doc_lang_en")],
                [InlineKeyboardButton("◀️ " + t(lang, "back"), callback_data="menu_create")],
            ]),
            parse_mode=ParseMode.MARKDOWN,
        )
        return

        # Циклограмма использует язык и подтверждённые данные профиля сразу,
        # поэтому не спрашивает их повторно и не проходит общий опросник.
        if doc_type == KINDERGARTEN_CYCLE_SCHEDULE:
            await self._start_cycle_schedule(query, context, user, lang)
            return
        if doc_type == DEVELOPMENT_MONITORING:
            await self._start_development_monitoring(query, context, user, lang)
            return

        # Для документов по ученику/ребёнку — предложить выбрать из базы
        if doc_type in STUDENT_LINKED_DOC_TYPES:
            students = await self.db.get_students(user_id)
            if students:
                keyboard = []
                for s in students[:8]:
                    keyboard.append([InlineKeyboardButton(
                        f"👤 {s['name']} ({s['class_name']})",
                        callback_data=f"gen_student_{s['id']}"
                    )])
                keyboard.append([InlineKeyboardButton(
                    "✍️ Ввести вручную" if lang == "ru" else "✍️ Қолмен енгізу",
                    callback_data="gen_student_0"
                )])
                context.user_data["doc_type"] = doc_type
                is_kg = user.get("role") == "kindergarten"
                title = ("👥 *Выберите воспитанника из базы:*" if is_kg else "👥 *Выберите ученика из базы:*") if lang == "ru" \
                    else ("👥 *Базадан баланы таңдаңыз:*" if is_kg else "👥 *Базадан оқушыны таңдаңыз:*")
                await query.edit_message_text(title, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode=ParseMode.MARKDOWN)
                return

        context.user_data["doc_type"]    = doc_type
        context.user_data["doc_answers"] = {}
        context.user_data["q_index"]     = 0

        saved_doc_lang = user.get("document_lang") or user.get("doc_language")
        if saved_doc_lang in {"ru", "kz", "en"}:
            context.user_data["doc_lang"] = saved_doc_lang
            language_label = _doc_lang_name(saved_doc_lang)
            await query.edit_message_text(
                (f"Язык документа: {language_label}" if saved_doc_lang == "ru" else f"Құжат тілі: {language_label}" if saved_doc_lang == "kz" else f"Document language: {language_label}"),
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("Создать" if saved_doc_lang == "ru" else "Жасау" if saved_doc_lang == "kz" else "Create", callback_data="doc_lang_confirm")],
                    [InlineKeyboardButton("Изменить язык" if saved_doc_lang == "ru" else "Тілді өзгерту" if saved_doc_lang == "kz" else "Change language", callback_data="doc_lang_change")],
                ])
            )
            return

        # ── Спросить язык документа ──
        doc_name = DOC_NAMES.get(lang, DOC_NAMES["ru"]).get(doc_type, doc_type)
        text = (
            f"📄 *{doc_name}*\n"
            f"{DIVIDER}\n"
            f"🌐 На каком языке создать документ?"
        ) if lang == "ru" else (
            f"📄 *{doc_name}*\n"
            f"{DIVIDER}\n"
            f"🌐 Құжатты қандай тілде жасау керек?"
        )
        keyboard = [
            [
                InlineKeyboardButton("🇷🇺 Русский", callback_data="doc_lang_ru"),
                InlineKeyboardButton("🇰🇿 Қазақша", callback_data="doc_lang_kz"),
            ],
            [InlineKeyboardButton("🇬🇧 English",    callback_data="doc_lang_en")],
            [InlineKeyboardButton("◀️ " + t(lang, "back"), callback_data="menu_create")],
        ]
        await query.edit_message_text(text, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode=ParseMode.MARKDOWN)

    async def _start_cycle_schedule(self, query, context, user, lang):
        lang = context.user_data.get("doc_lang", lang)
        answers = {
            "organization": user.get("school", ""),
            "educator_name": user.get("name", ""),
            "lang": context.user_data.get("doc_lang", user.get("lang", lang)),
        }
        # age_group в существующем профиле сада содержит название группы/возраст.
        if user.get("age_group"):
            answers["group"] = user["age_group"]

        questions = []
        if not answers.get("group"):
            questions.append({"key": "group", "q": "Для какой группы сделать циклограмму?" if lang == "ru" else "Циклограмма қай топқа керек?"})
        questions.extend([
            {"key": "period", "q": "На какую неделю или даты нужна циклограмма?" if lang == "ru" else "Қай аптаға немесе қай күндерге керек?"},
            {"key": "week_topic", "q": "Какая тема недели? Например: Домашние животные, Лето, Транспорт" if lang == "ru" else "Аптаның тақырыбы қандай? Мысалы: Үй жануарлары, Жаз, Көлік"},
            {"key": "events", "q": "Есть обязательные занятия, праздники или мероприятия на этой неделе? Если нет, напиши: нет" if lang == "ru" else "Осы аптада міндетті сабақтар, мерекелер немесе іс-шаралар бар ма? Егер жоқ болса, «жоқ» деп жазыңыз"},
        ])

        context.user_data["doc_type"] = KINDERGARTEN_CYCLE_SCHEDULE
        context.user_data["doc_lang"] = user.get("document_lang") or user.get("doc_language") or lang
        context.user_data["doc_answers"] = answers
        context.user_data["questions"] = questions
        context.user_data["q_index"] = 0
        context.user_data["step"] = "waiting_answer"
        doc_name = DOC_NAMES.get(context.user_data["doc_lang"], DOC_NAMES["ru"])[KINDERGARTEN_CYCLE_SCHEDULE]
        await self._ask_question(query.message, context, lang, 0, edit=True, query=query, doc_name=doc_name)

    async def _start_development_monitoring(self, query, context, user, lang):
        questions = [
            {"key": "period", "q": "За какой период нужен мониторинг?" if lang == "ru" else "Мониторинг қай кезеңге керек?"},
            {"key": "group", "q": "Для какой группы он нужен?" if lang == "ru" else "Ол қай топқа арналған?"},
            {"key": "children", "q": "Пришли список детей или загрузи таблицу с именами" if lang == "ru" else "Балалардың тізімін немесе аттары бар кестені жіберіңіз"},
        ]
        context.user_data.update({"doc_type": DEVELOPMENT_MONITORING, "doc_lang": lang, "doc_answers": {"organization": user.get("school", ""), "educator_name": user.get("name", ""), "age_group": user.get("age_group", "")}, "questions": questions, "q_index": 0, "step": "waiting_answer"})
        await self._ask_question(query.message, context, lang, 0, edit=True, query=query, doc_name=DOC_NAMES.get(lang, DOC_NAMES["ru"])[DEVELOPMENT_MONITORING])

    async def _use_student_data(self, query, context, user_id, user, lang, student_id):
        doc_type = context.user_data.get("doc_type", "")
        context.user_data["doc_answers"] = {}
        context.user_data["q_index"]     = 0

        if student_id > 0:
            student = await self.db.get_student(student_id)
            if student:
                grades       = json.loads(student.get("grades", "{}"))
                achievements = json.loads(student.get("achievements", "[]"))
                grades_str   = ", ".join(f"{k}-{v}" for k, v in grades.items()) if grades else "нет данных"
                achieve_str  = ", ".join(achievements) if achievements else "нет"
                context.user_data["doc_answers"]["student_name"] = f"{student['name']}, {student['class_name']}"
                context.user_data["doc_answers"]["performance"]  = grades_str
                context.user_data["doc_answers"]["activities"]   = achieve_str
                context.user_data["doc_answers"]["behavior"]     = student.get("behavior", "хорошее")

        saved_doc_lang = user.get("document_lang") or user.get("doc_language")
        if saved_doc_lang in {"ru", "kz", "en"}:
            context.user_data["doc_lang"] = saved_doc_lang
            q_lang = saved_doc_lang if saved_doc_lang in DOC_QUESTIONS else "ru"
            context.user_data["questions"] = REGISTRY_QUESTIONS.get(saved_doc_lang, {}).get(doc_type) or DOC_QUESTIONS.get(q_lang, DOC_QUESTIONS["ru"]).get(doc_type, [])
            context.user_data["step"] = "waiting_answer"
            await self._ask_question(query.message, context, lang, 0, edit=True, query=query, doc_name=DOC_NAMES.get(saved_doc_lang, DOC_NAMES["ru"]).get(doc_type, doc_type))
            return

        # Спросить язык документа
        doc_name = DOC_NAMES.get(lang, DOC_NAMES["ru"]).get(doc_type, doc_type)
        text = f"📄 *{doc_name}*\n{DIVIDER}\n🌐 На каком языке создать документ?" if lang == "ru" else f"📄 *{doc_name}*\n{DIVIDER}\n🌐 Құжатты қандай тілде жасау керек?"
        keyboard = [
            [InlineKeyboardButton("🇷🇺 Русский", callback_data="doc_lang_ru"),
             InlineKeyboardButton("🇰🇿 Қазақша", callback_data="doc_lang_kz")],
            [InlineKeyboardButton("🇬🇧 English",  callback_data="doc_lang_en")],
        ]
        await query.edit_message_text(text, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode=ParseMode.MARKDOWN)

    async def _ask_question(self, message, context, lang, idx, edit=False, query=None, doc_name=""):
        lang = context.user_data.get("doc_lang", lang)
        qs = context.user_data.get("questions", [])

        # Пропускаем вопросы у которых уже есть ответ
        while idx < len(qs) and qs[idx]["key"] in context.user_data.get("doc_answers", {}):
            idx += 1

        context.user_data["q_index"] = idx

        if idx >= len(qs):
            if edit and query:
                await query.edit_message_text("⏳ Начинаю генерацию..." if lang == "ru" else "⏳ Жасалуда...")
            await self._generate(message, context, lang)
            return

        q      = qs[idx]
        total  = len(qs)
        bar    = _progress_bar(idx + 1, total)
        doc_lang = context.user_data.get("doc_lang", lang)

        text = (
            f"📄 *{doc_name}* {_doc_lang_name(doc_lang)}\n"
            f"{bar}\n"
            f"{DIVIDER}\n\n"
            f"{q['q']}"
        ) if doc_name else (
            f"{bar}\n{DIVIDER}\n\n{q['q']}"
        )

        keyboard = [
            [InlineKeyboardButton("❌ " + t(lang, "cancel"), callback_data="doc_cancel")],
        ]

        if edit and query:
            await query.edit_message_text(text, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode=ParseMode.MARKDOWN)
        else:
            await message.reply_text(text, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode=ParseMode.MARKDOWN)

    async def handle_text(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        user_id = update.effective_user.id
        user    = await self.db.get_user(user_id)
        lang    = user.get("lang", "ru") if user else "ru"
        text    = update.message.text.strip()
        step    = context.user_data.get("step", "")

        if len(text) < 1:
            if (step == "waiting_answer" and
                    context.user_data.get("doc_type") == KINDERGARTEN_CYCLE_SCHEDULE):
                qs = context.user_data.get("questions", [])
                idx = context.user_data.get("q_index", 0)
                if idx < len(qs) and qs[idx]["key"] in CYCLE_REQUIRED_FIELDS:
                    await update.message.reply_text(f"Заполните поле: {CYCLE_REQUIRED_FIELDS[qs[idx]['key']]}.")
                    return
            await update.message.reply_text(t(lang, "val_too_short"))
            return
        if len(text) > 500:
            await update.message.reply_text(t(lang, "val_too_long"))
            return

        if step == "waiting_answer":
            qs  = context.user_data.get("questions", [])
            idx = context.user_data.get("q_index", 0)
            if idx < len(qs):
                if qs[idx]["key"] == "week_topic" and text.strip().lower() in {"план", "не знаю", "младшая группа", "неделя младшей группы", "білмеймін", "жоспар"}:
                    await update.message.reply_text("Уточните тему недели: например, «Домашние животные» или «Транспорт»." if lang == "ru" else "Апта тақырыбын нақтылаңыз: мысалы, «Үй жануарлары» немесе «Көлік».")
                    return
                context.user_data["doc_answers"][qs[idx]["key"]] = text
            doc_type = context.user_data.get("doc_type", "")
            doc_lang = context.user_data.get("doc_lang", lang)
            doc_name = DOC_NAMES.get(doc_lang, DOC_NAMES["ru"]).get(doc_type, "")
            if doc_type == DEVELOPMENT_MONITORING and idx + 1 >= len(qs):
                await update.message.reply_text("У тебя есть реальные наблюдения и уровни развития по детям?" if lang == "ru" else "Балалар бойынша нақты бақылаулар мен даму деңгейлері бар ма?", reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("Да, отправлю данные" if lang == "ru" else "Иә, деректерді жіберемін", callback_data="doc_monitoring_data_yes")], [InlineKeyboardButton("Нет, нужен пустой бланк" if lang == "ru" else "Жоқ, бос бланк керек", callback_data="doc_monitoring_data_no")], [InlineKeyboardButton("Отмена" if lang == "ru" else "Бас тарту", callback_data="doc_cancel")]]))
            else:
                await self._ask_question(update.message, context, lang, idx + 1, doc_name=doc_name)

        elif step == "monitoring_data":
            rows = []
            for line in text.splitlines():
                cells = [c.strip() for c in line.split("|")]
                if cells and cells[0]: rows.append(cells[:7])
            context.user_data["doc_answers"]["rows"] = rows
            await self._generate_monitoring(update.message, context, lang)

        elif step == "help_write":
            await self._handle_help_write(update, context, user, lang, text)

    async def handle_document(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        if context.user_data.get("step") != "template_upload":
            return
        user = await self.db.get_user(update.effective_user.id)
        lang = user.get("lang", "ru") if user else "ru"
        if not user or user.get("role") != "kindergarten":
            context.user_data.clear()
            await update.message.reply_text(
                "📄 Загрузка образцов доступна только воспитателям детского сада."
                if lang == "ru" else
                "📄 Үлгі жүктеу тек балабақша тәрбиешілеріне қолжетімді."
            )
            return
        document = update.message.document
        name = document.file_name or "template.docx"
        if not name.lower().endswith(".docx"):
            await update.message.reply_text("Нужен Word-файл в формате .docx. Отправь образец еще раз" if lang == "ru" else "Word .docx форматындағы файл қажет. Үлгіні қайта жіберіңіз")
            return
        folder = os.path.join(os.path.dirname(os.path.dirname(__file__)), "user_templates")
        os.makedirs(folder, exist_ok=True)
        dtype = context.user_data["template_doc_type"]
        path = os.path.join(folder, f"{update.effective_user.id}_{dtype}.docx")
        await (await document.get_file()).download_to_drive(custom_path=path)
        context.user_data["template_pending"] = {"doc_type": dtype, "path": path, "name": name, "metadata": {"telegram_file_id": document.file_id}}
        await update.message.reply_text(
            "Как использовать этот образец?" if lang == "ru" else "Бұл үлгіні қалай қолдану керек?",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("Только для меня" if lang == "ru" else "Тек мен үшін", callback_data="doc_template_personal")], [InlineKeyboardButton("Для всей организации" if lang == "ru" else "Бүкіл ұйым үшін", callback_data="doc_template_org")], [InlineKeyboardButton("Отмена" if lang == "ru" else "Бас тарту", callback_data="doc_cancel")]])
        )

    async def _generate_monitoring(self, message, context, lang):
        from handlers.word_generator import generate_word
        answers = context.user_data.get("doc_answers", {})
        children = [x.strip() for x in answers.get("children", "").replace("\n", ",").split(",") if x.strip()]
        rows = answers.get("rows", [])
        filename = generate_word("", DOC_NAMES.get(lang, DOC_NAMES["ru"])[DEVELOPMENT_MONITORING], monitoring_data={**answers, "children": children, "rows": rows, "lang": lang})
        with open(filename, "rb") as f:
            await message.reply_document(document=f, filename=f"monitoring_{datetime.now().strftime('%d%m%Y')}.docx", caption="📄 Мониторинг развития" if lang == "ru" else "📄 Даму мониторингі")
        os.remove(filename)
        user = await self.db.get_user(message.chat_id)
        await self.db.save_document(message.chat_id, DEVELOPMENT_MONITORING, DOC_NAMES.get(lang, DOC_NAMES["ru"])[DEVELOPMENT_MONITORING], "", 100)
        await self.db.log_analytics(message.chat_id, DEVELOPMENT_MONITORING, 100, lang)
        if user and not user.get("subscribed"):
            await self.db.increment_free(message.chat_id)
        context.user_data.clear()

    async def _handle_help_write(self, update, context, user, lang, user_input):
        field    = context.user_data.get("help_field", "")
        doc_type = context.user_data.get("doc_type", "")
        doc_lang = context.user_data.get("doc_lang", lang)
        doc_name = DOC_NAMES.get(doc_lang, DOC_NAMES["ru"]).get(doc_type, "")
        is_kg    = user.get("role") == "kindergarten"

        if is_kg:
            context_line = f"Возрастная группа: {user.get('age_group')}."
            asker = "Воспитатель"
        else:
            context_line = f"Предмет: {user.get('subject')}, классы: {user.get('classes')}."
            asker = "Учитель"

        client = anthropic.Anthropic(api_key=self.api_key)
        prompt = (
            f"{asker} просит помочь составить текст для поля «{field}» документа «{doc_name}».\n"
            f"{context_line}\n"
            f"Что хочет сказать: {user_input}\n\n"
            f"Предложи 2 коротких варианта официального текста для вставки в документ."
        )
        msg = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}]
        )
        result = msg.content[0].text
        context.user_data["step"] = "waiting_answer"
        await update.message.reply_text(
            f"✍️ *Варианты текста:*\n\n{result}\n\n{DIVIDER}\n_Скопируйте нужный вариант и отправьте_",
            parse_mode=ParseMode.MARKDOWN
        )

    async def _generate(self, message, context, lang):
        user_id  = message.chat_id
        user     = await self.db.get_user(user_id)
        doc_type = context.user_data.get("doc_type", "")
        doc_lang = context.user_data.get("doc_lang", lang)
        answers  = context.user_data.get("doc_answers", {})
        doc_name = DOC_NAMES.get(doc_lang, DOC_NAMES["ru"]).get(doc_type, doc_type)

        if doc_type == KINDERGARTEN_CYCLE_SCHEDULE:
            missing = validate_cycle_schedule_answers(answers)
            if missing:
                await message.reply_text("Не могу создать циклограмму. Заполните поле: " + ", ".join(missing) + ".")
                return

        if doc_type == DEVELOPMENT_MONITORING:
            await self._generate_monitoring(message, context, lang)
            return

        # Красивое сообщение о генерации
        gen_msgs = {
            "ru": f"⚙️ *Генерирую {doc_name}...*\n\n🌐 Язык: {_doc_lang_name(doc_lang)}\n⭐ Качество: автоматическая проверка\n\n_Напиши /cancel чтобы отменить_",
            "kz": f"⚙️ *{doc_name} жасалуда...*\n\n🌐 Тіл: {_doc_lang_name(doc_lang)}\n⭐ Сапа: автоматты тексеру\n\n_Болдырмау үшін /cancel жазыңыз_",
            "en": f"⚙️ *Generating {doc_name}...*\n\n🌐 Language: {_doc_lang_name(doc_lang)}\n⭐ Quality: auto-check\n\n_Type /cancel to stop_",
        }
        await message.reply_text(gen_msgs.get(lang, gen_msgs["ru"]), parse_mode=ParseMode.MARKDOWN)

        answers_text  = "\n".join(f"- {k}: {v}" for k, v in answers.items())

        # Профиль пользователя — подставляем автоматически
        profile_ctx = _build_profile_context(user, doc_lang)

        # Расписание из агентной памяти (если загружено) — только если реально есть данные
        schedule_ctx = ""
        try:
            from handlers.agent import AgentHandler
            schedule_ctx = await AgentHandler(self.db, self.api_key).get_context_for_generation(user_id)
        except Exception:
            pass

        system_prompt = get_system_prompt(user, doc_lang)

        # Образцы, загруженные админом для этого типа документа — добавляем как эталон
        samples_ctx = ""
        try:
            samples = await self.db.get_samples_for_type(doc_type, doc_lang, limit=1)
            if not samples:
                # Образец может быть на другом языке: используем его только как
                # каркас, а текст результата всё равно создаём на doc_lang.
                samples = [s for s in await self.db.get_all_samples(limit=200) if s["doc_type"] == doc_type and s["is_active"]][:1]
            if samples:
                samples_ctx = "\nЭТАЛОННЫЙ ОБРАЗЕЦ (используй только структуру, порядок разделов и формат таблиц; текст переведи на выбранный язык):\n" + samples[0]["content"][:3000]
        except Exception:
            pass

        # Личный Word-образец воспитателя — прочитать и передать генератору,
        # иначе сохранённый файл никак не влияет на результат.
        try:
            personal_template = await self.db.get_user_template(user_id, doc_type)
            if personal_template and os.path.exists(personal_template["file_path"]):
                from docx import Document
                template_doc = Document(personal_template["file_path"])
                template_parts = [p.text.strip() for p in template_doc.paragraphs if p.text.strip()]
                for table in template_doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):
                            template_parts.append(" | ".join(cells))
                template_text = "\n".join(template_parts)[:5000]
                if template_text:
                    samples_ctx += (
                        "\nЛИЧНЫЙ ОБРАЗЕЦ ПОЛЬЗОВАТЕЛЯ (сохрани его структуру, "
                        "названия блоков и формат таблиц):\n" + template_text
                    )
        except Exception as exc:
            print(f"Template read error: {exc}")

        language_instruction = {
            "ru": "Пиши весь документ только на русском языке.",
            "kz": "Құжаттың барлық мәтінін тек қазақ тілінде жаз. Орысша үлгі берілсе де, оны аударып, тек құрылымы мен кесте қаңқасын сақта.",
            "en": "Write the entire document only in English. If the sample is in another language, translate its text while preserving only its structure and table layout.",
        }.get(doc_lang, "Пиши весь документ только на русском языке.")
        user_prompt = (
            f"{language_instruction}\n"
            f"Создай документ: {doc_name}\n\n"
            f"{profile_ctx}\n"
            f"{schedule_ctx}\n"
            f"{samples_ctx}\n\n"
            f"Данные для этого документа:\n{answers_text}\n\n"
            f"Используй профиль автоматически. Если данных не хватает — пиши [уточнить], не выдумывай."
        )

        if doc_type == KINDERGARTEN_CYCLE_SCHEDULE:
            user_prompt += (
                "\n\nЭто циклограмма для детского сада. Не выдумывай даты, ФИО,"
                " организацию, режим дня или мероприятия. Составь только нейтральные"
                " планируемые идеи по указанной теме; мероприятия используй только если"
                " пользователь указал их явно."
            )

        client  = anthropic.Anthropic(api_key=self.api_key)
        result  = ""
        score   = 0
        attempts = 0

        while score < 85 and attempts < 2:
            attempts += 1
            if attempts > 1:
                improve_msg = {
                    "ru": "🔄 *Улучшаю качество...*\n_Аз қалды!_",
                    "kz": "🔄 *Сапаны жақсартуда...*\n_Аз қалды!_",
                    "en": "🔄 *Improving...*\n_Almost done!_",
                }
                await message.reply_text(improve_msg.get(lang, improve_msg["ru"]), parse_mode=ParseMode.MARKDOWN)

            msg = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=3000,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}]
            )
            result = msg.content[0].text

            eval_msg = client.messages.create(
                model="claude-haiku-4-5",
                max_tokens=10,
                messages=[{"role": "user", "content": SELF_EVAL_PROMPT.format(document=result[:3000])}]
            )
            try:
                score = int("".join(filter(str.isdigit, eval_msg.content[0].text.strip()[:5])))
            except:
                score = 88

            if score < 85 and attempts < 2:
                user_prompt = (
                    f"Улучши документ. Оценка {score}/100. "
                    f"Заполни все пустые поля, добавь конкретику из профиля.\n\n{result}"
                )

        # Формируем Word документ
        wait_msg = await message.reply_text(
            "📄 Формирую документ..." if lang == "ru" else "📄 Құжат жасалуда..."
        )

        try:
            from handlers.word_generator import generate_word
            filename = generate_word(
                result,
                doc_name,
                user.get("name", ""),
                cycle_data=answers if doc_type == KINDERGARTEN_CYCLE_SCHEDULE else None,
                monitoring_data=answers if doc_type == "kg_individual_development_card" else None,
                registry_doc_type=doc_type,
            )
            with open(filename, "rb") as f:
                caption = {
                    "ru": f"📄 *{doc_name}*\n✅ Готов к печати",
                    "kz": f"📄 *{doc_name}*\n✅ Басуға дайын",
                }
                await message.reply_document(
                    document=f,
                    filename=f"{doc_name}_{datetime.now().strftime('%d%m%Y')}.docx",
                    caption=caption.get(lang, caption["ru"]),
                    parse_mode=ParseMode.MARKDOWN
                )
            os.remove(filename)
        except Exception as e:
            print(f"Word error: {e}")
            import traceback
            traceback.print_exc()
            for chunk in range(0, len(result), 4000):
                await message.reply_text(result[chunk:chunk+4000])

        try:
            await wait_msg.delete()
        except:
            pass

        # Сохранить в БД
        await self.db.save_document(user_id, doc_type, doc_name, result, score)
        await self.db.log_analytics(user_id, doc_type, score, doc_lang)

        # Обновляем счётчик бесплатных — берём свежие данные из БД
        fresh_user = await self.db.get_user(user_id)
        if not fresh_user.get("subscribed"):
            await self.db.increment_free(user_id)
            fresh_user = await self.db.get_user(user_id)
            free_used = fresh_user.get("free_used", 0)
            total_free = free_limit_for(fresh_user)
            free_left = max(0, total_free - free_used)

            if free_left == 0:
                from handlers.profile import TIER_PRICES
                promo_available = not fresh_user.get("promo_used")

                if promo_available:
                    kb_after = [
                        [InlineKeyboardButton(
                            f"🔥 PRO за {TIER_PRICES['pro_promo']} тг (вместо {TIER_PRICES['pro']})" if lang == "ru"
                            else f"🔥 PRO {TIER_PRICES['pro_promo']} тг",
                            callback_data="prof_choose_pro_promo"
                        )],
                        [InlineKeyboardButton("🎁 Или пригласить и получить +5", callback_data="menu_invite")],
                        [InlineKeyboardButton("🏠 Главное меню", callback_data="menu_main")],
                    ]
                    note = (
                        f"⚠️ *Бесплатные документы исчерпаны!*\n\n"
                        f"🔥 Месяц PRO с безлимитной генерацией — всего *{TIER_PRICES['pro_promo']} тг* "
                        f"(вместо {TIER_PRICES['pro']} тг). Разовое предложение для вас.\n"
                        f"Или пригласите коллегу через /invite — за каждого +5 документов."
                    ) if lang == "ru" else (
                        f"⚠️ *Тегін құжаттар таусылды!*\n\n"
                        f"🔥 Шексіз PRO — небары *{TIER_PRICES['pro_promo']} тг* ({TIER_PRICES['pro']} тг орнына).\n"
                        f"Немесе /invite арқылы әріптесіңізді шақырыңыз."
                    )
                else:
                    kb_after = [
                        [InlineKeyboardButton("🎁 Пригласить и получить +5", callback_data="menu_invite")],
                        [InlineKeyboardButton("⭐ Оформить PRO — безлимит", callback_data="prof_sub")],
                        [InlineKeyboardButton("🏠 Главное меню", callback_data="menu_main")],
                    ]
                    note = ("⚠️ *Бесплатные документы исчерпаны!*\n"
                            "Оформите PRO или пригласите коллегу через /invite — за каждого +5 документов.") if lang == "ru" else \
                           ("⚠️ *Тегін құжаттар таусылды!*\n"
                            "PRO рәсімдеңіз немесе /invite арқылы әріптесіңізді шақырыңыз.")
            else:
                kb_after = [
                    [InlineKeyboardButton("📄 Создать ещё", callback_data="menu_create")],
                    [InlineKeyboardButton("🏠 Главное меню", callback_data="menu_main")],
                ]
                note = (f"🆓 Осталось бесплатных: *{free_left}/{total_free}*") if lang == "ru" else \
                       (f"🆓 Қалған тегін: *{free_left}/{total_free}*")
        else:
            kb_after = [
                [InlineKeyboardButton("📄 Создать ещё", callback_data="menu_create")],
                [InlineKeyboardButton("🏠 Главное меню", callback_data="menu_main")],
            ]
            note = "⭐ *PRO* — безлимитный доступ" if lang == "ru" else "⭐ *PRO* — шексіз қол жеткізу"

        await message.reply_text(
            note,
            reply_markup=InlineKeyboardMarkup(kb_after),
            parse_mode=ParseMode.MARKDOWN
        )
        context.user_data.clear()
