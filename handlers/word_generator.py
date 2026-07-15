"""
Docura.kz — профессиональный Word генератор
Стандарт: Times New Roman 14pt, поля ГОСТ, кириллица
"""
import os
import re
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Цвета ─────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x2E, 0x5A)
BLUE  = RGBColor(0x2E, 0x75, 0xB6)
GRAY  = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x1A, 0x20, 0x2C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_HEX = "EBF3FB"
NAVY_HEX       = "1A2E5A"
BLUE_HEX       = "2E75B6"

FONT_MAIN = "Times New Roman"
FONT_SIZE_MAIN = 14  # ГОСТ РК для официальных документов


def _set_cell_color(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color_hex="C8D8E8", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tcPr.append(borders)


def _add_paragraph_border_left(para, color_hex=BLUE_HEX, size="18"):
    """Синяя полоска слева у параграфа"""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_hrule(doc, color_hex=BLUE_HEX, size="12"):
    """Горизонтальная линия"""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def _run(para, text, bold=False, italic=False, size=None, color=None, font=None):
    run = para.add_run(text)
    run.font.name  = font or FONT_MAIN
    run.font.size  = Pt(size or FONT_SIZE_MAIN)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT,
          bold=False, italic=False, size=None, color=None,
          space_before=0, space_after=4, first_indent=None,
          keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before    = Pt(space_before)
    p.paragraph_format.space_after     = Pt(space_after)
    p.paragraph_format.keep_with_next  = keep_with_next
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if text:
        _run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def validate_cycle_schedule_data(data: dict) -> list[str]:
    """Проверяет обязательные поля до создания циклограммы."""
    labels = {
        "group": "группа",
        "period": "неделя или даты",
        "week_topic": "тема недели",
    }
    return [label for key, label in labels.items() if not str(data.get(key, "")).strip()]


def generate_word(content: str, title: str, teacher_name: str = "", cycle_data: dict | None = None, monitoring_data: dict | None = None, registry_doc_type: str = "") -> str:
    """
    Генерирует красивый Word документ по стандартам РК.
    Возвращает путь к .docx файлу.
    """
    doc = Document()

    # ── Страница: поля ГОСТ РК ────────────────────────
    for section in doc.sections:
        section.left_margin   = Cm(3.0)   # 30мм
        section.right_margin  = Cm(1.5)   # 15мм
        section.top_margin    = Cm(2.0)   # 20мм
        section.bottom_margin = Cm(2.0)   # 20мм
        if cycle_data or monitoring_data:
            # Пять рабочих дней и пять блоков лучше читаются на альбомной странице.
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width    = Cm(29.7)
            section.page_height   = Cm(21.0)
        else:
            section.page_width    = Cm(21.0)  # A4
            section.page_height   = Cm(29.7)

    # ── Стиль Normal по умолчанию ─────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_MAIN
    style.font.size = Pt(FONT_SIZE_MAIN)
    style.paragraph_format.line_spacing = Pt(18)
    style.paragraph_format.space_after  = Pt(0)

    # ── ШАПКА — тёмно-синяя таблица ──────────────────
    hdr_table = doc.add_table(rows=1, cols=1)
    hdr_table.style = "Table Grid"
    hdr_cell = hdr_table.rows[0].cells[0]
    _set_cell_color(hdr_cell, NAVY_HEX)
    _set_cell_borders(hdr_cell, NAVY_HEX)
    hdr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hp = hdr_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(6)
    hp.paragraph_format.space_after  = Pt(6)
    hr = hp.add_run("Подготовлено в Docura")
    hr.font.name  = FONT_MAIN
    hr.font.size  = Pt(9)
    hr.font.color.rgb = WHITE
    hr.font.bold  = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── ЗАГОЛОВОК ДОКУМЕНТА ───────────────────────────
    title_p = _para(doc, title.upper(),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    bold=True, size=16, color=NAVY,
                    space_before=6, space_after=6)

    # Синяя линия под заголовком
    _add_hrule(doc, BLUE_HEX, "12")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── ПАРСИМ СОДЕРЖИМОЕ ─────────────────────────────
    if registry_doc_type == "kg_individual_development_card":
        _add_individual_development_card(doc, monitoring_data or {})
    elif cycle_data:
        _add_kindergarten_cycle_schedule(doc, cycle_data)
    elif monitoring_data:
        _add_development_monitoring(doc, monitoring_data)
    else:
        _parse_content(doc, content)

    if registry_doc_type != "kg_individual_development_card":
        doc.add_paragraph().paragraph_format.space_before = Pt(12)
        _add_hrule(doc, "A0AEC0", "6")
        _para(doc, "Подготовлено в Docura", align=WD_ALIGN_PARAGRAPH.CENTER,
              italic=True, size=8, color=RGBColor(0xA0, 0xAE, 0xC0), space_before=4)

    fname = os.path.join(tempfile.gettempdir(), f"docura_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(fname)
    return fname


def _add_individual_development_card(doc, data: dict):
    """Пустая карта: результаты не заполняются без подтверждённых наблюдений."""
    kz = data.get("lang") == "kz"
    labels = (["Баланың аты-жөні", "Туған жылы", "Баланың жасы", "Тобы"] if kz
              else ["Child’s full name", "Birth year", "Child’s age", "Group"] if data.get("lang") == "en"
              else ["ФИО ребенка", "Год рождения", "Возраст ребенка", "Группа"])
    values = [data.get("child_name", ""), data.get("birth_year", ""), data.get("age", ""), data.get("group", "")]
    info = doc.add_table(rows=4, cols=2); info.style = "Table Grid"
    for row, label, value in zip(info.rows, labels, values):
        _set_cell_color(row.cells[0], LIGHT_BLUE_HEX); _run(row.cells[0].paragraphs[0], label, bold=True, size=10); _run(row.cells[1].paragraphs[0], value, size=10)
    headers = (["Құзыреттіліктер", "Бастапқы бақылау", "Аралық бақылау", "Қорытынды бақылау", "Қорытынды, даму деңгейі"] if kz
               else ["Competencies", "Initial observation", "Intermediate observation", "Final observation", "Conclusion / level"] if data.get("lang") == "en"
               else ["Компетенции", "Начальное наблюдение", "Промежуточное наблюдение", "Итоговое наблюдение", "Вывод / уровень"])
    rows = ["Физикалық қасиеттерді дамыту", "Коммуникативтік дағдыларды дамыту", "Танымдық және зияткерлік дағдыларды дамыту", "Шығармашылық дағдыларды дамыту", "Әлеуметтік-эмоционалды дағдыларды қалыптастыру"] if kz else (["Physical development", "Communication skills", "Cognitive and intellectual skills", "Creative and research skills", "Social-emotional skills"] if data.get("lang") == "en" else ["Развитие физических качеств", "Развитие коммуникативных навыков", "Развитие познавательных навыков", "Развитие творческих навыков", "Формирование социально-эмоциональных навыков"])
    table = doc.add_table(rows=1 + len(rows), cols=5); table.style = "Table Grid"
    for i, header in enumerate(headers): _set_cell_color(table.rows[0].cells[i], NAVY_HEX); _run(table.rows[0].cells[i].paragraphs[0], header, bold=True, size=7, color=WHITE)
    for r, name in enumerate(rows, 1):
        _run(table.rows[r].cells[0].paragraphs[0], name, size=8)
    legend = ["I деңгей: төмен", "II деңгей: орташа", "III деңгей: жоғары"] if kz else (["Level I: low", "Level II: medium", "Level III: high"] if data.get("lang") == "en" else ["I уровень: низкий", "II уровень: средний", "III уровень: высокий"])
    _para(doc, "\n".join(legend), size=9, space_before=8)


def _add_kindergarten_cycle_schedule(doc, data: dict):
    """Создаёт циклограмму только из подтверждённых данных пользователя."""
    missing = validate_cycle_schedule_data(data)
    if missing:
        raise ValueError("Не заполнены обязательные поля: " + ", ".join(missing))

    def value(key):
        return str(data.get(key, "")).strip() or "________________"

    info = [
        ("Организация", value("organization")),
        ("Группа", value("group")),
        ("Период", value("period")),
        ("Воспитатель", value("educator_name")),
    ]
    info_table = doc.add_table(rows=len(info), cols=2)
    info_table.style = "Table Grid"
    for row, (label, text) in zip(info_table.rows, info):
        _set_cell_color(row.cells[0], LIGHT_BLUE_HEX)
        _set_cell_borders(row.cells[0])
        _set_cell_borders(row.cells[1])
        _run(row.cells[0].paragraphs[0], label + ":", bold=True, size=11, color=NAVY)
        _run(row.cells[1].paragraphs[0], text, size=11)

    _para(doc, "Планируемое содержание недели", bold=True, color=NAVY,
          space_before=10, space_after=5)
    topic = value("week_topic")
    events = str(data.get("events", "")).strip()
    has_events = events and events.lower() not in {"нет", "жоқ", "none", "no"}
    days = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница"]
    blocks = [
        "Утренний прием и гимнастика",
        "ОУД",
        "Прогулка",
        "Сон и закаливание",
        "Игры и уход детей домой",
    ]
    daily = [
        [f"Беседа по теме «{topic}»; гимнастика с наблюдениями.", f"ОУД: «Знакомство с темой {topic}». Цель: выделять признаки. Игра: «Найди пару».", f"Наблюдение по теме «{topic}»; подвижная игра «Повтори движение».", "Подготовка ко сну по режиму группы; дыхательное упражнение.", f"Игра «Назови по теме {topic}»; беседа при уходе."],
        [f"Рассматривание иллюстраций по теме «{topic}»; новая гимнастика.", f"ОУД: «Признаки и действия». Цель: расширять словарь. Игра: «Что изменилось?».", "Подвижная игра с правилами; наблюдение на прогулке.", "Расслабление перед сном; процедуры по режиму группы.", f"Творческая игра по теме «{topic}»; индивидуальная беседа."],
        [f"Утренний круг по теме «{topic}»; ритмическая гимнастика.", f"ОУД: «Сравни и выбери». Цель: учить сравнивать. Игра: «Лишний предмет».", "Поисковая прогулка; игра-эстафета без инвентаря.", "Чтение перед сном; процедуры по режиму группы.", "Настольная игра «Собери картинку»; итоги дня."],
        [f"Пальчиковая игра по теме «{topic}»; координационная гимнастика.", f"ОУД: «Где и как бывает». Цель: закреплять связи. Игра: «Кому что нужно?».", "Наблюдение и игра «Опиши, не называя».", "Спокойная музыка перед сном; процедуры по режиму группы.", f"Сюжетно-ролевая игра по теме «{topic}»; обмен впечатлениями."],
        [f"Повторение слов по теме «{topic}»; выбор движений детьми.", f"ОУД: «Итоги недели». Цель: применять знания. Игра: «Угадай по описанию».", "Игра-наблюдение «Что запомнили?»; двигательная активность.", "Подготовка ко сну по режиму группы; расслабление.", "Свободные игры и беседа об итогах темы."],
    ]
    if data.get("lang") == "kz":
        days = ["Дүйсенбі", "Сейсенбі", "Сәрсенбі", "Бейсенбі", "Жұма"]
        blocks = ["Таңертеңгі қабылдау және гимнастика", "ҰОҚ", "Серуен", "Ұйқы және шынықтыру", "Ойындар және үйге қайту"]
        daily = [
            [f"«{topic}» тақырыбы бойынша әңгіме; бақылау жаттығулары.", f"ҰОҚ: «{topic} тақырыбымен танысу». Мақсат: белгілерін ажырату. Ойын: «Жұбын тап».", f"«{topic}» бойынша бақылау; «Қимылды қайтала» ойыны.", "Топ режимі бойынша ұйқыға дайындық; тыныс алу жаттығуы.", f"«{topic}» бойынша атау ойыны; ата-анамен қысқа әңгіме."],
            [f"«{topic}» суреттерін қарау; жаңа гимнастика.", "ҰОҚ: «Белгілері мен әрекеттері». Мақсат: сөздік қорын кеңейту. Ойын: «Не өзгерді?».", "Ережелі қимылды ойын; серуендегі бақылау.", "Ұйқы алдындағы босаңсу; топ режимі бойынша процедуралар.", f"«{topic}» бойынша шығармашылық ойын; жеке әңгіме."],
            [f"«{topic}» бойынша таңғы шеңбер; ырғақты гимнастика.", "ҰОҚ: «Салыстыр және таңда». Мақсат: салыстыруға үйрету. Ойын: «Артық зат».", "Іздеу серуені; құралсыз эстафета ойыны.", "Ұйқы алдындағы оқу; топ режимі бойынша процедуралар.", "Үстел ойыны; күн қорытындысы."],
            [f"«{topic}» бойынша саусақ ойыны; үйлестіру гимнастикасы.", "ҰОҚ: «Қайда және қалай болады». Мақсат: байланыстарды бекіту. Ойын: «Кімге не керек?».", "Бақылау және «Атын атамай сипатта» ойыны.", "Ұйқы алдындағы баяу музыка; топ режимі бойынша процедуралар.", f"«{topic}» бойынша сюжеттік ойын; әсерлермен алмасу."],
            [f"«{topic}» сөздерін қайталау; балалар қимылды таңдайды.", "ҰОҚ: «Апта қорытындысы». Мақсат: білімді ойында қолдану. Ойын: «Сипаттама бойынша тап».", "«Не есте қалды?» бақылау ойыны; қимыл белсенділігі.", "Топ режимі бойынша ұйқыға дайындық; босаңсу.", "Еркін ойындар және тақырып қорытындысы."],
        ]

    table = doc.add_table(rows=1 + len(blocks), cols=1 + len(days))
    table.style = "Table Grid"
    headers = ["Режимный момент", *days]
    for c_idx, header in enumerate(headers):
        cell = table.rows[0].cells[c_idx]
        _set_cell_color(cell, NAVY_HEX)
        _set_cell_borders(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, header, bold=True, size=9, color=WHITE)

    for r_idx, block in enumerate(blocks, start=1):
        label_cell = table.rows[r_idx].cells[0]
        _set_cell_color(label_cell, LIGHT_BLUE_HEX)
        _set_cell_borders(label_cell)
        _run(label_cell.paragraphs[0], block, bold=True, size=9, color=NAVY)
        for c_idx in range(1, len(days) + 1):
            cell = table.rows[r_idx].cells[c_idx]
            _set_cell_borders(cell)
            text = daily[c_idx - 1][r_idx - 1]
            # Пользовательское мероприятие отражается как план, а не как факт.
            if has_events and r_idx == 2:
                text += f" Учесть при планировании: {events}."
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _run(p, text, size=8)



def _add_development_monitoring(doc, data: dict):
    lang = data.get("lang", "ru")
    labels = ["Организация", "Группа", "Возраст", "Период", "Воспитатель"] if lang == "ru" else ["Ұйым", "Топ", "Жасы", "Кезең", "Тәрбиеші"]
    values = [data.get("organization") or "________________", data.get("group") or "________________", data.get("age_group") or "________________", data.get("period") or "________________", data.get("educator_name") or "________________"]
    info = doc.add_table(rows=len(labels), cols=2); info.style = "Table Grid"
    for row, label, value in zip(info.rows, labels, values):
        _set_cell_color(row.cells[0], LIGHT_BLUE_HEX); _run(row.cells[0].paragraphs[0], label + ":", bold=True, size=10); _run(row.cells[1].paragraphs[0], value, size=10)
    headers = ["№", "ФИО ребенка", "Физическое развитие", "Коммуникативное развитие", "Познавательное развитие", "Творческое развитие", "Социально-эмоциональное развитие", "Примечание"] if lang == "ru" else ["№", "Баланың аты-жөні", "Дене дамуы", "Коммуникативтік даму", "Танымдық даму", "Шығармашылық даму", "Әлеуметтік-эмоциялық даму", "Ескертпе"]
    rows = data.get("rows") or [[name] for name in data.get("children", [])] or [[] for _ in range(10)]
    table = doc.add_table(rows=1 + len(rows), cols=8); table.style = "Table Grid"
    for i, h in enumerate(headers): _set_cell_color(table.rows[0].cells[i], NAVY_HEX); _run(table.rows[0].cells[i].paragraphs[0], h, bold=True, size=7, color=WHITE)
    for n, source in enumerate(rows, 1):
        vals = [str(n)] + (source + [""] * 7)[:7]
        for i, val in enumerate(vals): _run(table.rows[n].cells[i].paragraphs[0], val, size=8)


def _parse_content(doc, content: str):
    lines = content.split("\n")

    def is_section(s):
        if not s or len(s) < 3: return False
        if re.match(r'^(\d+\.|[IVX]+\.)\s+', s): return True
        no_punct = re.sub(r'[^a-zA-Zа-яёА-ЯЁіІқҚғҒүҮұҰіәӘһҺ]', '', s)
        return (len(no_punct) > 2
                and no_punct == no_punct.upper()
                and no_punct.upper() != no_punct.lower())

    def is_sub(s):
        return (s.endswith(":") and 4 < len(s) < 80
                and not s.startswith(("-","•","*","–")))

    def is_bullet(s):
        return s.startswith(("- ","• ","* ","– ","· "))

    def is_sig(s):
        kw = ["подпись","директор","учитель","кл.рук","дата:","м.п.",
              "классный руководитель","қолы","мұғалім","күні:",
              "мектеп директоры","________","___/"]
        return any(k in s.lower() for k in kw)

    def is_table(s):
        return s.count("|") >= 2

    prev_empty = False
    i = 0
    while i < len(lines):
        s = lines[i].strip()

        if not s:
            if not prev_empty:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
            prev_empty = True
            i += 1
            continue
        prev_empty = False

        # Таблица Markdown
        if is_table(s):
            tbl_lines = []
            while i < len(lines) and (is_table(lines[i]) or
                  re.match(r'^\s*\|[\s\-:]+\|', lines[i])):
                raw = lines[i].strip()
                if not re.match(r'^\|[\s\-:]+\|', raw):
                    tbl_lines.append(raw)
                i += 1
            if tbl_lines:
                _add_table(doc, tbl_lines)
            continue

        # Заголовок раздела — синяя плашка + полоска
        if is_section(s):
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            p = doc.add_paragraph()
            p.paragraph_format.space_before   = Pt(8)
            p.paragraph_format.space_after    = Pt(4)
            p.paragraph_format.left_indent    = Cm(0.3)
            p.paragraph_format.keep_with_next = True
            _add_paragraph_border_left(p, BLUE_HEX, "20")
            # Светло-голубой фон через shading параграфа
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  LIGHT_BLUE_HEX)
            pPr.append(shd)
            r = p.add_run("  " + s)
            r.font.name      = FONT_MAIN
            r.font.size      = Pt(FONT_SIZE_MAIN)
            r.font.bold      = True
            r.font.color.rgb = NAVY
            i += 1
            continue

        # Подзаголовок
        if is_sub(s):
            p = _para(doc, s, bold=True, color=GRAY,
                      space_before=6, space_after=3,
                      keep_with_next=True)
            i += 1
            continue

        # Буллет
        if is_bullet(s):
            text = re.sub(r'^[-•*–·]\s+', '', s)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after        = Pt(2)
            bullet_run = p.add_run("▸  ")
            bullet_run.font.name      = "Arial"
            bullet_run.font.size      = Pt(11)
            bullet_run.font.color.rgb = BLUE
            text_run = p.add_run(text)
            text_run.font.name      = FONT_MAIN
            text_run.font.size      = Pt(FONT_SIZE_MAIN)
            text_run.font.color.rgb = BLACK
            i += 1
            continue

        # Подпись
        if is_sig(s):
            p = _para(doc, s, color=GRAY,
                      space_before=3, space_after=3)
            i += 1
            continue

        # Обычный абзац
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after        = Pt(3)
        r = p.add_run(s)
        r.font.name      = FONT_MAIN
        r.font.size      = Pt(FONT_SIZE_MAIN)
        r.font.color.rgb = BLACK
        i += 1


def _add_table(doc, lines: list):
    """Красивая таблица с синей шапкой и чередующимися строками"""
    rows_data = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    max_cols = max(len(r) for r in rows_data)
    rows_data = [r + [""] * (max_cols - len(r)) for r in rows_data]

    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = "Table Grid"

    # Ширина колонок
    page_width_cm = 21.0 - 3.0 - 1.5  # A4 - поля
    col_width = Cm(page_width_cm / max_cols)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    for r_idx, row_data in enumerate(rows_data):
        is_header = (r_idx == 0)
        is_even   = (r_idx % 2 == 0)
        bg = NAVY_HEX if is_header else ("EBF3FB" if is_even else "FFFFFF")

        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            _set_cell_color(cell, bg)
            _set_cell_borders(cell, "C8D8E8" if not is_header else NAVY_HEX)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Очищаем дефолтный параграф
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)

            r = p.add_run(cell_text)
            r.font.name      = FONT_MAIN
            r.font.size      = Pt(11)
            r.font.bold      = is_header
            r.font.color.rgb = WHITE if is_header else BLACK

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
